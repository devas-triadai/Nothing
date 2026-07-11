import re
import logging
from pathlib import Path
from typing import List, Optional
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from api.models.compliance_models import ClauseResultData, PipelineResult

logger = logging.getLogger("agra.compliance_docx")


def _sanitize_filename(name: str) -> str:
    safe = re.sub(r'[\\/:*?"<>|]', '_', name)
    return safe[:100].strip()


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def _verdict_color(verdict: Optional[str]) -> str:
    if not verdict:
        return "D3D3D3"
    v = verdict.upper()
    if v == "COMPLIANT":
        return "22C55E"
    elif v == "PARTIAL":
        return "EAB308"
    elif v == "NON_COMPLIANT":
        return "EF4444"
    elif v == "UNVERIFIABLE":
        return "9CA3AF"
    return "D3D3D3"


def _verdict_text_color(verdict: Optional[str]) -> str:
    if not verdict:
        return "000000"
    v = verdict.upper()
    if v in ("COMPLIANT", "NON_COMPLIANT"):
        return "FFFFFF"
    return "000000"


def generate_compliance_report(
    result: PipelineResult,
    reference_name: str,
    output_path: str,
    selected_standard_names: Optional[List[str]] = None,
) -> str:
    doc = DocxDocument()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOTR Compliance Verification Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(reference_name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4a, 0x8b, 0xff)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated: {__import__('datetime').datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    total = result.total_clauses
    scored = result.compliant_count + result.partial_count + result.non_compliant_count
    denom = scored if scored > 0 else total
    score_pct = (result.compliant_count + result.partial_count * 0.5) / denom * 100 if denom > 0 else 0

    doc.add_paragraph(
        f"This report evaluates a vendor submission against the SOTR requirements "
        f"for the project: {reference_name}. "
        f"Of {total} clauses evaluated, {result.compliant_count} are COMPLIANT, "
        f"{result.partial_count} are PARTIAL, {result.non_compliant_count} are NON_COMPLIANT, "
        f"and {result.unverifiable_count} are UNVERIFIABLE. "
        f"The overall compliance score is {score_pct:.1f}%."
    )

    rec = result.recommendation.value if result.recommendation else "N/A"
    rec_p = doc.add_paragraph()
    run = rec_p.add_run(f"Recommendation: {rec}")
    run.bold = True
    run.font.size = Pt(13)

    if result.missing_clause_count > 0:
        doc.add_paragraph(f"Missing clauses detected: {result.missing_clause_count}")
    if result.contradiction_count > 0:
        doc.add_paragraph(f"Contradictions detected: {result.contradiction_count}")
    if result.house_rule_violation_count > 0:
        doc.add_paragraph(f"House rule violations detected: {result.house_rule_violation_count}")

    doc.add_paragraph("")

    # ── Clause-by-Clause Table ──
    doc.add_heading("Clause-by-Clause Evaluation", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    headers = ["Clause ID", "Source", "Verdict", "Severity", "Finding"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for clause in result.clauses:
        row_cells = table.add_row().cells
        row_cells[0].text = clause.clause_id
        row_cells[1].text = clause.source_file
        row_cells[2].text = clause.verdict.value if clause.verdict else "—"
        row_cells[3].text = clause.severity.value if clause.severity else "—"
        row_cells[4].text = clause.finding[:200] if clause.finding else "—"

        v = clause.verdict.value if clause.verdict else ""
        color = _verdict_color(v)
        text_color = _verdict_text_color(v)
        _set_cell_shading(row_cells[2], color)
        for p in row_cells[2].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(
                    int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16)
                )

        for cell in row_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph("")

    # ── Non-Compliance Register ──
    non_comp = [c for c in result.clauses if c.verdict and c.verdict.value in ("NON_COMPLIANT", "PARTIAL")]
    if non_comp:
        doc.add_heading("Non-Compliance Register", level=1)
        table2 = doc.add_table(rows=1, cols=5)
        table2.style = 'Light Grid Accent 1'
        hdr2 = table2.rows[0].cells
        for i, h in enumerate(["Clause ID", "Verdict", "Severity", "Recommendation", "Finding"]):
            hdr2[i].text = h
            for p in hdr2[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for clause in non_comp:
            row_cells = table2.add_row().cells
            row_cells[0].text = clause.clause_id
            row_cells[1].text = clause.verdict.value if clause.verdict else "—"
            row_cells[2].text = clause.severity.value if clause.severity else "—"
            row_cells[3].text = clause.recommendation.value if clause.recommendation else "—"
            row_cells[4].text = clause.finding[:200] if clause.finding else "—"

            v = clause.verdict.value if clause.verdict else ""
            color = _verdict_color(v)
            _set_cell_shading(row_cells[1], color)

            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

        doc.add_paragraph("")

    # ── House Rule Deviations ──
    hr_violations = [c for c in result.clauses if c.house_rule_flag and c.house_rule_flag.violated]
    if hr_violations:
        doc.add_heading("House Rule Deviations", level=1)
        doc.add_paragraph("The following clauses have house rule / standards violations:")
        table3 = doc.add_table(rows=1, cols=4)
        table3.style = 'Light Grid Accent 1'
        hdr3 = table3.rows[0].cells
        for i, h in enumerate(["Clause ID", "Rule Reference", "Verdict", "Note"]):
            hdr3[i].text = h
            for p in hdr3[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for clause in hr_violations:
            row_cells = table3.add_row().cells
            row_cells[0].text = clause.clause_id
            row_cells[1].text = clause.house_rule_flag.rule_reference
            row_cells[2].text = clause.verdict.value if clause.verdict else "—"
            row_cells[3].text = clause.house_rule_flag.note
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

        doc.add_paragraph("")

    # ── Contradictions ──
    contradictions = [c for c in result.clauses if c.contradictions]
    if contradictions:
        doc.add_heading("Contradictions Detected", level=1)
        for clause in contradictions:
            for ct in clause.contradictions:
                p = doc.add_paragraph()
                run = p.add_run(f"{clause.clause_id}: ")
                run.bold = True
                doc.add_paragraph(f"  Between: {', '.join(ct.between)}", style='List Bullet')
                doc.add_paragraph(f"  Statement A: {ct.statement_a}", style='List Bullet')
                doc.add_paragraph(f"  Statement B: {ct.statement_b}", style='List Bullet')
                if ct.note:
                    doc.add_paragraph(f"  Note: {ct.note}", style='List Bullet')

    # ── Missing Clauses ──
    missing = [c for c in result.clauses if c.is_missing]
    if missing:
        doc.add_heading("Missing Clauses", level=1)
        doc.add_paragraph("The following SOTR clauses have no corresponding vendor evidence:")
        for clause in missing:
            p = doc.add_paragraph(f"{clause.clause_id}: {clause.requirement_text[:150]}", style='List Bullet')

    # ── Historical Feedback ──
    has_history = any(c.historical_notes for c in result.clauses)
    if has_history:
        doc.add_heading("Historical Feedback", level=1)
        for clause in result.clauses:
            if clause.historical_notes:
                for note in clause.historical_notes:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{clause.clause_id} (Run #{note.run_id}): ")
                    run.bold = True
                    trend_str = f" [{note.trend.upper()}]" if note.trend else ""
                    doc.add_paragraph(f"  {note.note}{trend_str}", style='List Bullet')

    # ── Standards Reference Appendix ──
    if selected_standard_names:
        doc.add_heading("Appendix: Standards Referenced", level=1)
        for s in selected_standard_names:
            doc.add_paragraph(s, style='List Bullet')

    # ── Footer Watermark ──
    try:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "AI-Generated Draft - ICG AGRA"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    except Exception as e:
        logger.warning("Could not add footer: %s", e)

    doc.save(output_path)
    logger.info("Compliance report saved to %s", output_path)
    return output_path
