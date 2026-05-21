"""
AGRA Phase 2 — Router: Compliance Check Engine (v2)
Ground-up rewrite. Key fixes vs v1:
  - raw=True on ALL llm_engine.generate() calls for JSON — prevents clean_llm_output()
    from stripping lines containing "topic", "Task:", "Role:" etc.
  - Correct _sanitize_json_content() — state-machine, only touches inside strings
  - Single-pass architecture — no multi-batch complexity, no streaming
  - 3 JSON repair layers max (was 6)
  - Smart standard recommendation via cosine similarity
  - Always returns a response — never hangs
"""

import json
import logging
import re
import time
import uuid
from pathlib import Path
from typing import List, Optional

import numpy as np
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from docx import Document as DocxDocument

from api.utils.auth_check import get_current_user
from api.utils.usage_logger import log_usage
from api.rag import embedder, llm as llm_engine
from api.rag.vector_store import get_store

logger = logging.getLogger("agra.compliance")

router = APIRouter()

import os as _os
_DATA_DIR = Path(_os.environ.get("AGRA_DATA_DIR", "/workspace/agra_data"))
if not _DATA_DIR.exists():
    _DATA_DIR = Path(__file__).resolve().parent.parent.parent / "agra_data"
_OUTPUTS_DIR = _DATA_DIR / "outputs"
_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

_CORS_HEADERS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "POST, GET, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type, Authorization",
}

# ─────────────────────────────────────────────────────────────────────────────
# JSON HELPERS  (3 layers max — no more)
# ─────────────────────────────────────────────────────────────────────────────

def _clean_json(raw: str) -> str:
    """Strip markdown fences and extract outermost [...] array."""
    s = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.MULTILINE)
    s = re.sub(r"```\s*$", "", s.strip(), flags=re.MULTILINE)
    s = s.strip()
    start = s.find("[")
    end = s.rfind("]") + 1
    if start != -1 and end > start:
        return s[start:end]
    return s


def _sanitize_json_content(s: str) -> str:
    """
    Fix raw control characters INSIDE JSON string values only.
    State-machine approach — never touches structural quotes or backslashes.
    Raw newline/tab/CR inside a string value → escaped \\n / \\t / \\r.
    Other control chars (< 0x20) inside string → removed.
    """
    result = []
    in_string = False
    i = 0
    while i < len(s):
        c = s[i]
        if in_string:
            if c == "\\":
                result.append(c)
                if i + 1 < len(s):
                    i += 1
                    result.append(s[i])
            elif c == '"':
                in_string = False
                result.append(c)
            elif c == "\n":
                result.append("\\n")
            elif c == "\r":
                result.append("\\r")
            elif c == "\t":
                result.append("\\t")
            elif ord(c) < 32:
                pass  # drop other control chars inside strings
            else:
                result.append(c)
        else:
            if c == '"':
                in_string = True
                result.append(c)
            else:
                result.append(c)
        i += 1
    return "".join(result)


def _repair_truncated_json(s: str) -> str:
    """Close open braces/brackets and remove trailing commas to fix truncated JSON."""
    s = s.strip()
    if not s:
        return "[]"
    # Close unterminated string
    in_string = False
    escaped = False
    for c in s:
        if escaped:
            escaped = False
            continue
        if c == "\\":
            escaped = True
            continue
        if c == '"':
            in_string = not in_string
    if in_string:
        s += '"'
    # Strip trailing comma before closing
    s = re.sub(r",\s*$", "", s)
    # Count and close open structures
    open_braces = s.count("{") - s.count("}")
    open_brackets = s.count("[") - s.count("]")
    s += "}" * max(0, open_braces)
    s += "]" * max(0, open_brackets)
    # Remove trailing commas before closing brackets/braces
    s = re.sub(r",\s*([}\]])", r"\1", s)
    return s


def _extract_json_objects_regex(raw: str) -> list:
    """Last resort — extract individual {...} objects from garbled output."""
    findings = []
    pattern = r"\{[^{}]*(?:\{[^{}]*\}[^{}]*)?\}"
    for match in re.findall(pattern, raw, re.DOTALL):
        if not any(k in match.lower() for k in ["clause", "verdict", "finding", "requirement"]):
            continue
        try:
            obj = json.loads(_sanitize_json_content(match))
            if isinstance(obj, dict):
                findings.append(obj)
        except Exception:
            continue
    return findings


def _parse_llm_json(raw: str) -> list:
    """
    Parse LLM output into a list of finding dicts.
    3 layers: clean → sanitize → parse → repair → regex.
    Always returns a list (never raises).
    """
    # Layer 1: clean + sanitize + parse
    try:
        cleaned = _clean_json(raw)
        sanitized = _sanitize_json_content(cleaned)
        data = json.loads(sanitized)
        if isinstance(data, list):
            return [d for d in data if isinstance(d, dict)]
        if isinstance(data, dict):
            return [data]
    except (json.JSONDecodeError, ValueError):
        pass

    # Layer 2: repair truncated + retry
    try:
        cleaned = _clean_json(raw)
        sanitized = _sanitize_json_content(cleaned)
        repaired = _repair_truncated_json(sanitized)
        data = json.loads(repaired)
        if isinstance(data, list):
            logger.info("JSON repair succeeded: %d items", len(data))
            return [d for d in data if isinstance(d, dict)]
    except Exception:
        pass

    # Layer 3: regex object extraction
    findings = _extract_json_objects_regex(raw)
    if findings:
        logger.info("Regex extraction recovered %d findings", len(findings))
        return findings

    logger.warning("All JSON parse strategies failed. Raw (first 300): %s", repr(raw[:300]))
    return []


# ─────────────────────────────────────────────────────────────────────────────
# COSINE SIMILARITY HELPER
# ─────────────────────────────────────────────────────────────────────────────

def _cosine_sim(a: list, b: list) -> float:
    """Cosine similarity between two embedding vectors."""
    va = np.array(a, dtype=np.float32)
    vb = np.array(b, dtype=np.float32)
    denom = np.linalg.norm(va) * np.linalg.norm(vb)
    if denom == 0:
        return 0.0
    return float(np.dot(va, vb) / denom)


def _mean_embedding(chunks: list) -> list:
    """Compute mean embedding across a list of document chunks."""
    texts = [c["text"] for c in chunks if c.get("text", "").strip()][:10]
    if not texts:
        return []
    vecs = embedder.embed_texts(texts)
    arr = np.array(vecs, dtype=np.float32)
    mean = arr.mean(axis=0)
    norm = np.linalg.norm(mean)
    if norm > 0:
        mean = mean / norm
    return mean.tolist()


def _select_best_chunks(chunks: list, query_embedding: list, max_chars: int) -> str:
    """
    Select the most relevant chunks up to max_chars total.
    Scores each chunk by cosine similarity to the query embedding,
    returns the top chunks concatenated as a string.
    """
    if not chunks:
        return ""
    if not query_embedding:
        # Fallback: just take first N chars
        text = "\n\n".join(c["text"] for c in chunks)
        return text[:max_chars]

    scored = []
    for c in chunks:
        text = c.get("text", "").strip()
        if not text:
            continue
        vec = embedder.embed_query(text[:500])  # embed first 500 chars for speed
        score = _cosine_sim(query_embedding, vec)
        scored.append((score, c))

    scored.sort(key=lambda x: x[0], reverse=True)

    result_parts = []
    total = 0
    for _, c in scored:
        fname = c.get("metadata", {}).get("filename", "")
        text = c.get("text", "")
        part = f"[{fname}]: {text}" if fname else text
        if total + len(part) > max_chars:
            remaining = max_chars - total
            if remaining > 100:
                result_parts.append(part[:remaining])
            break
        result_parts.append(part)
        total += len(part)

    return "\n\n".join(result_parts)


# ─────────────────────────────────────────────────────────────────────────────
# PYDANTIC MODELS
# ─────────────────────────────────────────────────────────────────────────────

class RecommendRequest(BaseModel):
    subject_doc_id: str = Field(..., description="Doc ID of the subject document")


class ComplianceRequest(BaseModel):
    subject_doc_ids: List[str] = Field(..., min_length=1)
    standard_doc_ids: List[str] = Field(..., min_length=1)
    check_scope: Optional[str] = Field(None)


# ─────────────────────────────────────────────────────────────────────────────
# CORS OPTIONS HANDLERS
# ─────────────────────────────────────────────────────────────────────────────

@router.options("/compliance/recommend-standards")
async def recommend_options():
    return JSONResponse(content={}, headers=_CORS_HEADERS)


@router.options("/compliance/check")
async def check_options():
    return JSONResponse(content={}, headers=_CORS_HEADERS)


# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINT 1: RECOMMEND STANDARDS
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/compliance/recommend-standards")
async def recommend_standards(
    body: RecommendRequest,
    user: dict = Depends(get_current_user),
):
    """
    Given a subject document ID, returns all available standard documents
    ranked by cosine similarity to the subject doc's content.
    Standards with similarity >= 0.30 are flagged as recommended.
    """
    store = get_store()

    # Load subject chunks
    subject_chunks = store.get_chunks_by_doc(body.subject_doc_id)
    if not subject_chunks:
        raise HTTPException(
            status_code=404,
            detail="Subject document not found in knowledge base.",
            headers=_CORS_HEADERS,
        )

    # Compute subject mean embedding
    subject_embedding = _mean_embedding(subject_chunks)

    # Get all documents from store
    all_docs = store.list_unique_documents()
    standard_docs = [
        d for d in all_docs
        if d.get("doc_id") != body.subject_doc_id
        and (
            d.get("doc_id", "").startswith("builtin:")
            or (d.get("category", "").lower() in ["standard", "global standard", "sotr", "imo", "rule"])
            or (d.get("document_type", "") == "standard")
        )
    ]

    if not standard_docs:
        # If no docs are tagged as standards, return all non-subject docs ranked
        standard_docs = [d for d in all_docs if d.get("doc_id") != body.subject_doc_id]

    recommendations = []
    for std_doc in standard_docs:
        std_chunks = store.get_chunks_by_doc(std_doc["doc_id"])
        if not std_chunks:
            continue

        std_embedding = _mean_embedding(std_chunks)
        score = _cosine_sim(subject_embedding, std_embedding) if subject_embedding and std_embedding else 0.0

        # Generate a one-line reason using LLM (only for top candidates to save tokens)
        reason = ""
        if score >= 0.25:
            try:
                subject_snippet = subject_chunks[0]["text"][:400] if subject_chunks else ""
                std_snippet = std_chunks[0]["text"][:400] if std_chunks else ""
                reason_prompt = (
                    f"Subject document excerpt:\n{subject_snippet}\n\n"
                    f"Standard document excerpt:\n{std_snippet}\n\n"
                    "In one sentence (max 20 words), explain why this standard is relevant to the subject document. "
                    "Be specific about which aspect matches. Output the sentence only."
                )
                reason = llm_engine.generate(
                    [{"role": "user", "content": reason_prompt}],
                    max_tokens=60,
                    temperature=0.1,
                    raw=True,
                ).strip()
                # Clean up any markdown or extra whitespace
                reason = re.sub(r"^[*_`#\-]+|[*_`#\-]+$", "", reason).strip()
            except Exception as e:
                logger.debug("Reason generation failed for %s: %s", std_doc["doc_id"], e)
                reason = ""

        recommendations.append({
            "doc_id": std_doc["doc_id"],
            "filename": std_doc.get("filename", "Unknown"),
            "category": std_doc.get("category", ""),
            "chunks": std_doc.get("chunks", 0),
            "score": round(score, 3),
            "recommended": score >= 0.30,
            "reason": reason,
        })

    # Sort by score descending
    recommendations.sort(key=lambda x: x["score"], reverse=True)
    logger.info(
        "Standard recommendation for %s: %d total, %d recommended",
        body.subject_doc_id,
        len(recommendations),
        sum(1 for r in recommendations if r["recommended"]),
    )

    return JSONResponse(
        content={"recommendations": recommendations},
        headers=_CORS_HEADERS,
    )


# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINT 2: COMPLIANCE CHECK (single-pass, non-streaming, raw=True)
# ─────────────────────────────────────────────────────────────────────────────

_ERROR_FINDING = {
    "topic": "Analysis Error",
    "clause_id": "N/A",
    "requirement": "Automated compliance parsing",
    "acceptance_criterion": "Valid structured output from LLM",
    "verdict": "Unverifiable",
    "severity": "Major",
    "finding": (
        "The compliance engine could not produce a structured analysis. "
        "This may be due to document complexity or LLM output format issues. "
        "Please try again with a narrower check scope or fewer documents."
    ),
    "recommendation": "Re-run with check_scope targeting specific sections, or reduce document count.",
    "citation": "N/A",
}


@router.post("/compliance/check")
async def compliance_check(
    request: Request,
    body: ComplianceRequest,
    user: dict = Depends(get_current_user),
):
    """
    Run clause-by-clause compliance analysis.
    Returns a JSON response with all findings + a .docx report download URL.

    Key design decisions vs v1:
    - raw=True on all LLM calls → clean_llm_output() never strips JSON keys
    - Single-pass: one LLM call, no multi-batch complexity
    - Cosine similarity to select best chunks (not just first N)
    - 3-layer JSON parsing max
    - Always returns — never hangs
    """
    start_time = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""

    store = get_store()

    # ── Load subject chunks ──
    subject_chunks = []
    for s_id in body.subject_doc_ids:
        subject_chunks.extend(store.get_chunks_by_doc(s_id))

    if not subject_chunks:
        raise HTTPException(
            status_code=404,
            detail="Subject documents not found in knowledge base.",
            headers=_CORS_HEADERS,
        )

    # ── Load standard chunks ──
    standard_chunks = []
    for std_id in body.standard_doc_ids:
        std = store.get_chunks_by_doc(std_id)
        if not std:
            raise HTTPException(
                status_code=404,
                detail=f"Standard document {std_id} not found.",
                headers=_CORS_HEADERS,
            )
        standard_chunks.extend(std)

    if not standard_chunks:
        raise HTTPException(
            status_code=400,
            detail="No standard document content found.",
            headers=_CORS_HEADERS,
        )

    subject_filenames = list({
        c["metadata"].get("filename", "Subject")
        for c in subject_chunks if "metadata" in c
    })
    subject_filenames_str = ", ".join(subject_filenames)
    scope_note = f"\nFocus specifically on: {body.check_scope}" if body.check_scope else ""

    # ── Select best chunks using cosine similarity ──
    # Token budget: 3328 total slot - 200 prompt template - 1200 output = 1928 input tokens
    # At ~4 chars/token: ~7700 chars. Split 55/45 subject/standard.
    _MAX_SUBJECT_CHARS = 2000
    _MAX_STANDARD_CHARS = 1800

    subject_query = f"compliance requirements {body.check_scope or ''} {subject_filenames_str}"
    try:
        query_embedding = embedder.embed_query(subject_query)
        subject_text = _select_best_chunks(subject_chunks, query_embedding, _MAX_SUBJECT_CHARS)
        standard_text = _select_best_chunks(standard_chunks, query_embedding, _MAX_STANDARD_CHARS)
    except Exception as e:
        logger.warning("Embedding-based chunk selection failed (%s), falling back to first-N", e)
        subject_text = "\n\n".join(
            f"[{c['metadata'].get('filename','Unknown')}]: {c['text']}"
            for c in subject_chunks[:8]
        )[:_MAX_SUBJECT_CHARS]
        standard_text = "\n\n".join(c["text"] for c in standard_chunks[:6])[:_MAX_STANDARD_CHARS]

    # ── Build prompt ──
    # Prompt ends with `[` so Gemma's first token continues the JSON array.
    # This prevents echo of the instruction header.
    prompt = (
        f"SUBJECT DOCUMENTS ({subject_filenames_str}):\n"
        f"{subject_text}\n\n"
        f"STANDARDS:\n"
        f"{standard_text}"
        f"{scope_note}\n\n"
        "Compare each standard clause against the subject documents.\n"
        "For each clause output one JSON object with these exact keys:\n"
        "  topic, clause_id, requirement, acceptance_criterion, verdict, severity, finding, recommendation, citation\n"
        "verdict must be one of: Compliant | Non-Compliant | Partial | Missing | Contradiction | Unverifiable\n"
        "severity must be one of: Critical | Major | Minor | None\n"
        "Output 5 to 8 findings. Return a JSON array only. No prose. No markdown fences.\n"
        "["
    )

    # ── LLM call — raw=True is critical ──
    findings_raw = "[]"
    try:
        findings_raw = llm_engine.generate(
            [{"role": "user", "content": prompt}],
            max_tokens=1200,
            temperature=0.1,
            raw=True,  # CRITICAL: bypass clean_llm_output() which strips JSON keys
        )
        logger.info("LLM compliance call completed, raw output length: %d", len(findings_raw))
    except Exception as e:
        logger.error("Compliance LLM call failed: %s", e)

    # The prompt ends with `[` and llama-server returns the continuation,
    # so we prepend `[` to reconstruct the full array.
    if findings_raw and not findings_raw.strip().startswith("["):
        findings_raw = "[" + findings_raw

    # ── Parse JSON (3 layers) ──
    findings = _parse_llm_json(findings_raw)

    # ── Fallback error finding ──
    if not findings:
        logger.error("All JSON parse strategies failed — returning error finding")
        findings = [_ERROR_FINDING]

    # ── OCR confidence post-processing ──
    conf_scores = [
        c["metadata"].get("ocr_confidence", 1.0)
        for c in subject_chunks if "metadata" in c
    ]
    avg_conf = sum(conf_scores) / len(conf_scores) if conf_scores else 1.0
    if avg_conf < 0.65:
        logger.warning("Low OCR confidence: %.2f — marking findings as Unverifiable", avg_conf)
        for f in findings:
            if f.get("verdict") != "Missing":
                f["verdict"] = "Unverifiable"
                f["finding"] = f"[LOW OCR CONFIDENCE: {avg_conf:.2f}] " + f.get("finding", "")

    # ── Compute summary statistics ──
    compliant = sum(1 for f in findings if f.get("verdict") == "Compliant")
    non_compliant = sum(1 for f in findings if f.get("verdict") == "Non-Compliant")
    partial = sum(1 for f in findings if f.get("verdict") == "Partial")
    missing = sum(1 for f in findings if f.get("verdict") == "Missing")
    contradiction = sum(1 for f in findings if f.get("verdict") == "Contradiction")
    unverifiable = sum(1 for f in findings if f.get("verdict") == "Unverifiable")

    total_evaluable = len(findings) - unverifiable
    compliance_score = round((compliant / total_evaluable * 100) if total_evaluable > 0 else 0.0, 1)

    critical_issues = sum(
        1 for f in findings
        if f.get("severity") == "Critical" and f.get("verdict") != "Compliant"
    )

    if critical_issues > 0 or compliance_score < 70:
        recommendation = "REJECT"
    elif non_compliant > 0 or missing > 0 or partial > 0:
        recommendation = "APPROVE WITH CONDITIONS (REVISE)"
    else:
        recommendation = "APPROVE"

    # ── Build DOCX report ──
    job_id = str(uuid.uuid4())
    docx_path = _OUTPUTS_DIR / f"{job_id}_compliance.docx"
    try:
        doc = DocxDocument()
        doc.add_heading("Compliance Analysis Report", level=1)
        doc.add_heading(f"Subjects: {subject_filenames_str}", level=2)
        doc.add_paragraph(f"Scope: {body.check_scope or 'Full Document'}")
        doc.add_paragraph("")
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(f"Overall Compliance Score: {compliance_score}%")
        rec_para = doc.add_paragraph(f"Final Recommendation: {recommendation}")
        rec_para.runs[0].bold = True
        doc.add_paragraph(f"Critical Deficiencies Found: {critical_issues}")
        doc.add_paragraph("")
        doc.add_heading("Summary Statistics", level=2)
        tbl = doc.add_table(rows=7, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("Total Clauses Checked", str(len(findings))),
            ("Compliant", str(compliant)),
            ("Non-Compliant", str(non_compliant)),
            ("Partial", str(partial)),
            ("Missing", str(missing)),
            ("Contradiction", str(contradiction)),
            ("Unverifiable", str(unverifiable)),
        ]
        for i, (label, val) in enumerate(rows_data):
            tbl.rows[i].cells[0].text = label
            tbl.rows[i].cells[1].text = val
        doc.add_paragraph("")
        doc.add_heading("Detailed Findings Register", level=2)
        for i, f in enumerate(findings, 1):
            doc.add_heading(f"Finding {i}: {f.get('clause_id', 'N/A')}", level=3)
            doc.add_paragraph(f"Verdict: {f.get('verdict', 'N/A')} | Severity: {f.get('severity', 'None')}")
            doc.add_paragraph(f"Topic: {f.get('topic', 'N/A')}")
            doc.add_paragraph(f"Requirement: {f.get('requirement', 'N/A')}")
            doc.add_paragraph(f"Acceptance Criterion: {f.get('acceptance_criterion', 'N/A')}")
            doc.add_paragraph(f"Finding: {f.get('finding', 'N/A')}")
            doc.add_paragraph(f"Recommendation: {f.get('recommendation', 'N/A')}")
            doc.add_paragraph(f"Citation: {f.get('citation', 'N/A')}")
            doc.add_paragraph("")
        doc.save(str(docx_path))
        download_url = f"/api/agent/download/{job_id}_compliance.docx"
        logger.info("DOCX report saved: %s", docx_path)
    except Exception as e:
        logger.error("DOCX generation failed: %s", e)
        download_url = None

    # ── Log usage ──
    elapsed_ms = (time.time() - start_time) * 1000
    log_usage(
        action_type="compliance",
        module="compliance",
        token=auth_tok,
        response_time_ms=elapsed_ms,
    )

    logger.info(
        "Compliance check done: %d findings, score=%.1f%%, rec=%s, elapsed=%.0fms",
        len(findings), compliance_score, recommendation, elapsed_ms,
    )

    return JSONResponse(
        content={
            "findings": findings,
            "summary": {
                "total": len(findings),
                "compliant": compliant,
                "non_compliant": non_compliant,
                "partial": partial,
                "missing": missing,
                "contradiction": contradiction,
                "unverifiable": unverifiable,
                "score": compliance_score,
                "recommendation": recommendation,
            },
            "download_url": download_url,
        },
        headers=_CORS_HEADERS,
    )


