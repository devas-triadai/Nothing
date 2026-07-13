"""
AGRA Phase 2 — Router: Content Generation (PPT, Summary, Quiz)
"""

import asyncio
import json
import logging
import re
import time
import uuid
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel, Field

from api.utils.auth_check import get_current_user
from api.utils.usage_logger import log_usage
from api.utils.genealogy_client import (
    check_superseded_status,
    get_document_lineage,
    format_superseded_warning,
    should_include_genealogy,
)
from api.rag import embedder, llm as llm_engine
from api.rag.vector_store import get_store
from api.rag.reranker import rerank
from api.generators.ppt_gen import build_pptx
from api.generators.pdf_gen import generate_summary_pdf, generate_quiz_pdf

# ── Async PPT job store (in-memory, per-process) ──
# Maps job_id -> {"status": "pending"|"done"|"error", "filename": str, "slides_json": str, "error": str}
_ppt_jobs: dict = {}
from docx import Document as DocxDocument

logger = logging.getLogger("agra.generate")

router = APIRouter()

import os as _os
_DATA_DIR = Path(_os.environ.get("AGRA_DATA_DIR", "/workspace/agra_data"))
if not _DATA_DIR.exists():
    _DATA_DIR = Path(__file__).resolve().parent.parent.parent / "agra_data"
_OUTPUTS_DIR = _DATA_DIR / "outputs"
_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)


def _extract_balanced(text: str, open_char: str, close_char: str) -> Optional[str]:
    """
    Depth-counting bracket matcher. Returns the first balanced substring
    starting with `open_char`, accounting for strings (single/double quotes)
    and escaped chars inside strings. Returns None if no balanced region found.
    """
    start = text.find(open_char)
    if start < 0:
        return None

    depth = 0
    in_string = False
    string_char = ""
    i = start
    n = len(text)

    while i < n:
        ch = text[i]
        if in_string:
            if ch == "\\" and i + 1 < n:
                # Skip escaped character
                i += 2
                continue
            if ch == string_char:
                in_string = False
        else:
            if ch in ('"', "'"):
                in_string = True
                string_char = ch
            elif ch == open_char:
                depth += 1
            elif ch == close_char:
                depth -= 1
                if depth == 0:
                    return text[start:i + 1]
        i += 1

    return None  # unbalanced


def _repair_truncated_array(raw: str) -> Optional[str]:
    """
    Attempt to recover a JSON array that was truncated by the token limit.
    Strategy: find the last complete object in the array and close the array.
    Returns a valid JSON array string, or None if recovery is not possible.
    """
    start = raw.find('[')
    if start < 0:
        return None

    text = raw[start:]

    # Find last complete slide object by scanning for last '}'
    last_close = text.rfind('}')
    if last_close < 0:
        return None

    # Trim everything after the last complete '}' and close the array
    candidate = text[:last_close + 1].rstrip(', \t\n') + ']'

    try:
        result = json.loads(candidate)
        if isinstance(result, list) and len(result) > 0:
            return candidate
    except json.JSONDecodeError:
        pass

    return None


def _extract_section_title(text_snippet: str) -> Optional[str]:
    """Extract a meaningful section title from a text snippet (sentence or heading)."""
    if not text_snippet:
        return None
    # Clean up the text
    text = text_snippet.strip().replace('\n', ' ').replace('  ', ' ')
    # Remove common prefixes like bullet markers
    text = re.sub(r'^[-*•]\s*', '', text)
    # Try to find a colon-separated heading (e.g., "System Architecture: Overview")
    if ':' in text and len(text.split(':')[0]) > 5:
        candidate = text.split(':')[0].strip()
        if 10 <= len(candidate) <= 50:
            return candidate
    # Try to find ALL CAPS headings
    words = text.split()
    if len(words) >= 2:
        # Check for title case pattern (first word capitalized)
        if words[0][0].isupper() if words[0] else False:
            # Take first 4-6 words as title
            title_words = words[:min(6, len(words))]
            candidate = ' '.join(title_words)
            # Remove trailing punctuation
            candidate = re.sub(r'[;:,]$', '', candidate)
            if 10 <= len(candidate) <= 55:
                return candidate
    return None


def _clean_json(raw: str) -> str:
    """
    Strip markdown code fences and extract the first balanced JSON array or
    object from LLM output. Uses depth-counting bracket matching so it does
    not get confused by strings containing brackets or trailing prose.
    """
    if not raw:
        return ""

    # 1. Strip markdown fences aggressively (anywhere in text)
    cleaned = re.sub(r'```(?:json)?\s*', '', raw, flags=re.IGNORECASE)
    cleaned = re.sub(r'```', '', cleaned)
    cleaned = cleaned.strip()

    # 2. Fix trailing commas before } or ] (common LLM JSON mistake)
    cleaned = re.sub(r',\s*([}\]])', r'\1', cleaned)

    # 3. Locate first '[' or '{' and pick whichever appears first
    arr_pos = cleaned.find('[')
    obj_pos = cleaned.find('{')

    if arr_pos < 0 and obj_pos < 0:
        return cleaned

    # Choose the one that appears earlier in the text
    use_array = (arr_pos >= 0) and (obj_pos < 0 or arr_pos < obj_pos)

    if use_array:
        result = _extract_balanced(cleaned, '[', ']')
        if result:
            return result

    result = _extract_balanced(cleaned, '{', '}')
    if result:
        return result

    return cleaned


def _repair_json_with_llm(broken: str, error_msg: str) -> Optional[dict]:
    """
    Last-ditch repair: ask the LLM to fix malformed JSON. Returns parsed
    dict on success, None if repair also fails.
    """
    try:
        repair_prompt = (
            "The following text was supposed to be valid JSON but failed to parse. "
            f"Parser error: {error_msg}\n\n"
            "Return ONLY the corrected JSON. No prose, no explanation, no markdown fences.\n\n"
            f"BROKEN JSON:\n{broken[:6000]}"
        )
        repaired_raw = llm_engine.generate(
            messages=[
                {"role": "user", "content": f"Fix the broken JSON below. Return ONLY valid JSON, nothing else.\n\n{repair_prompt}"},
            ],
            max_tokens=4096,
            temperature=0.0,
            response_format={"type": "json_object"},
            raw=True,
        )
        cleaned = _clean_json(repaired_raw)
        return json.loads(cleaned)
    except Exception as e:
        logger.warning("JSON repair pass failed: %s", e)
        return None


def _build_fallback_slides(topic: str, context_text: str, num_slides: int = 10) -> list:
    """
    Generate a structured fallback deck when the LLM fails to return valid JSON.
    Strips metadata artifacts, extracts section headings for slide titles, and
    adds light visual variety (section headers, two_column) where possible.
    """
    slides = [
        {"layout": "title", "title": topic, "subtitle": "Indian Coast Guard | AGRA Knowledge Management"},
    ]

    # Strip raw metadata artifacts emitted by _format_context before splitting
    _ARTIFACT_RE = re.compile(
        r'\[Document:[^\]]*\]|\[Page\s*\d+[^\]]*\]|\|\s*Page\s*\d+|\[Content truncated[^\]]*\]',
        re.IGNORECASE,
    )
    clean_context = _ARTIFACT_RE.sub('', context_text).strip()

    # Extract section headings (ALL-CAPS lines or Markdown headings) for slide titles
    heading_re = re.compile(r'^(?:#{1,3}\s+|[A-Z][A-Z0-9\s,.-]{4,60}$)', re.MULTILINE)
    headings = [h.strip().lstrip('#').strip() for h in heading_re.findall(clean_context) if len(h.strip()) > 4]

    # Split remaining content into sentences (filter very short/empty ones)
    sentences = [
        s.strip() for s in re.split(r'[\n]+|(?<=\.)\s+', clean_context)
        if len(s.strip()) > 25 and not _ARTIFACT_RE.search(s)
    ]

    content_slides = num_slides - 2  # exclude title + thank_you
    chunk_size = max(1, len(sentences) // max(1, content_slides))

    idx = 0
    for i in range(content_slides):
        chunk = sentences[idx:idx + chunk_size]
        idx += chunk_size
        # Use an extracted heading if available, otherwise derive from topic
        slide_title = headings[i] if i < len(headings) else f"{topic} — Section {i + 1}"
        slide_title = slide_title[:80]  # cap title length
        if not chunk:
            chunk = [f"Refer to source document for Section {i + 1} details."]
        bullets = [s[:200] for s in chunk[:6]]
        # Every 3rd content slide: use section_header layout for visual rhythm
        if i > 0 and i % 3 == 0 and content_slides > 4:
            slides.append({"layout": "section_header", "title": slide_title, "subtitle": ""})
        else:
            slides.append({"layout": "bullets", "title": slide_title, "bullets": bullets})

    slides.append({"layout": "thank_you", "title": "Thank You", "subtitle": "AGRA — AI-Powered Knowledge Management"})
    return slides


# ── Master LLM Prompt for Intelligent Slide Design ──
# NOTE: Gemma 4 IT does NOT reliably honour system-prompt JSON constraints,
# so ALL instructions live in the user message.
_PPT_SLIDE_LAYOUTS = """
LAYOUTS: title|section_header|bullets|two_column|table|diagram|chart|thank_you
bullets: {title, bullets[3-6]}
two_column: {title, left_column{header,items[]}, right_column{header,items[]}}
diagram: {title, diagram_data{type:"flowchart", nodes[{id,label,shape}], edges[{from,to}]}}
chart: {title, chart_data{type:"bar_chart", data{labels[],values[]}}}
table: {title, table_data{headers[], rows[[]]}}
"""


class PPTRequest(BaseModel):
    topic: str = Field(..., min_length=1)
    num_slides: int = Field(default=10, ge=3, le=25)
    doc_ids: List[str] = Field(default_factory=list)
    style_notes: Optional[str] = None
    # Revision / versioning fields
    revision_prompt: Optional[str] = Field(None, description="Instructions for revising an existing PPT")
    previous_slides_json: Optional[str] = Field(None, description="JSON of previous slides to revise")
    version: int = Field(default=1, ge=1, description="Version number (v1, v2, v3...)")


@router.get("/generate/ppt/status/{job_id}")
async def ppt_job_status(
    job_id: str,
    user: dict = Depends(get_current_user),
):
    """
    Poll the status of an async PPT generation job.
    Returns {status, download_url, filename, slides_json} when done.
    """
    job = _ppt_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="PPT job not found")
    if job["status"] == "done":
        return JSONResponse({
            "status": "done",
            "download_url": f"/api/agent/download/{job['filename']}",
            "filename": job["filename"],
            "slides_json": job.get("slides_json", ""),
        })
    if job["status"] == "error":
        return JSONResponse({"status": "error", "error": job.get("error", "Unknown error")})
    return JSONResponse({"status": "pending"})


async def _generate_slide_batch(prompt_text: str, max_tokens: int = 1500, temperature: float = 0.1, expected_count: int = 0) -> list:
    """Make one LLM call and return parsed slides list (empty list on failure).
    If expected_count is provided and we get fewer slides, retry once with stronger prompt."""
    for attempt in range(2):
        try:
            actual_prompt = prompt_text
            if attempt == 1 and expected_count > 0:
                # Retry with stronger emphasis
                actual_prompt = prompt_text.replace(
                    "Create exactly",
                    "IMPORTANT: You MUST create exactly"
                ) + f"\n\nCRITICAL: Return EXACTLY {expected_count} slide objects in the JSON array. No more, no less."

            raw = await asyncio.to_thread(
                llm_engine.generate, [{"role": "user", "content": actual_prompt}],
                max_tokens=max_tokens, temperature=temperature, raw=True,
            )
            if not raw or len(raw.strip()) < 10:
                continue
            cleaned = _clean_json(raw)
            if not cleaned:
                recovered = _repair_truncated_array(raw)
                cleaned = recovered or raw
            try:
                data = json.loads(cleaned)
            except json.JSONDecodeError:
                recovered = _repair_truncated_array(cleaned)
                if not recovered:
                    continue
                data = json.loads(recovered)
            if isinstance(data, list):
                slides = [d for d in data if isinstance(d, dict)]
                # If we got expected count or more, return immediately
                if not expected_count or len(slides) >= expected_count:
                    return slides
                # Otherwise retry if this was first attempt
                if attempt == 0:
                    logger.warning(f"Slide batch returned {len(slides)} slides, expected {expected_count}. Retrying...")
                    continue
                return slides
            if isinstance(data, dict):
                return [data]
        except Exception as e:
            logger.warning("Slide batch failed (attempt %d): %s", attempt + 1, e)
    return []


async def _generate_slides_multibatch(body: "PPTRequest", context_full: str, source_filenames: list) -> list:
    """
    Generate slides via 3-4 batched LLM calls for RICH content per slide.
    Splits document into thirds. Each call produces 3-4 detailed slides referencing
    the actual document content. Also generates a dedicated table slide and a
    flowchart/architecture slide.
    """
    topic = body.topic or "Document Overview"
    num_slides = body.num_slides

    # Split context into 3 slices to expose the FULL document across batches
    L = len(context_full)
    s1 = context_full[: L // 3][:4500]
    s2 = context_full[L // 3 : 2 * L // 3][:4500]
    s3 = context_full[2 * L // 3 :][:4500]

    # Filename hint for in-deck citations
    files_str = ", ".join(sorted({f for f in source_filenames if f})[:5]) or "the source document"

    # ── Batch 1: slides 1-3 (title + introduction + key concept) ──
    prompt1 = f"""DOCUMENT EXCERPT (Part 1 of 3, from {files_str}):
{s1}

Create exactly 3 PowerPoint slides about: {topic}
Slide 1: layout="title", title="{topic}", subtitle (presenter/org from doc)
Slide 2: layout="bullets", title="Overview" or similar, 4-5 detailed bullets summarizing the document's purpose, scope, stakeholders. Each bullet 12-25 words.
Slide 3: layout="bullets", title for main concept/problem statement, 4-5 detailed bullets with specifics from the document.

Return ONLY a JSON array of 3 objects. No prose. Start with [
["""
    batch1 = await _generate_slide_batch(prompt1, max_tokens=1800, temperature=0.1, expected_count=3)

    # ── Batch 2: slides 4-7 (technical/methodology with table) ──
    prompt2 = f"""DOCUMENT EXCERPT (Part 2 of 3, from {files_str}):
{s2}

Create exactly 4 PowerPoint slides continuing a presentation about: {topic}
Slide 1: layout="section_header", title for a major section in this excerpt
Slide 2: layout="bullets", 4-5 detailed bullets on methodology/approach/technical details, 12-25 words each
Slide 3: layout="table", title="Comparison" or feature/spec table, with table_data={{"headers":["Feature","Description","Status"], "rows":[[...],[...],[...]]}} (4-6 rows from the document)
Slide 4: layout="two_column", title, left_column=["..3-4 items.."], right_column=["..3-4 items.."] with real document content

Return ONLY a JSON array of 4 objects. No prose. Start with [
["""
    batch2 = await _generate_slide_batch(prompt2, max_tokens=2000, temperature=0.1, expected_count=4)

    # ── Batch 3: slides 8-10 (architecture flowchart + conclusion + thank_you) ──
    prompt3 = f"""DOCUMENT EXCERPT (Part 3 of 3, from {files_str}):
{s3}

Create exactly 3 PowerPoint slides concluding a presentation about: {topic}
Slide 1: layout="diagram", title="System Architecture" or similar, diagram_data={{"type":"flowchart","nodes":[{{"id":"A","label":"<step from doc>","shape":"rounded_rect"}},{{"id":"B","label":"<step>","shape":"rect"}},{{"id":"C","label":"<step>","shape":"rect"}},{{"id":"D","label":"<output>","shape":"rounded_rect"}}],"edges":[{{"from":"A","to":"B"}},{{"from":"B","to":"C"}},{{"from":"C","to":"D"}}]}} — use REAL terminology from the document for node labels
Slide 2: layout="bullets", title="Key Takeaways" or "Conclusion", 4-5 detailed bullets summarizing outcomes/benefits/next steps
Slide 3: layout="thank_you", title="Thank You", subtitle (organization name from doc, e.g. "Indian Coast Guard")

Return ONLY a JSON array of 3 objects. No prose. Start with [
["""
    batch3 = await _generate_slide_batch(prompt3, max_tokens=2000, temperature=0.1, expected_count=3)

    combined = []
    combined.extend(batch1)
    combined.extend(batch2)
    combined.extend(batch3)

    # If we got something reasonable, return it (may be 7-10 slides; pad/trim happens later)
    if len(combined) >= 3:
        return combined
    return []


async def _run_ppt_job(job_id: str, body: PPTRequest, auth_header_token: str):
    """Background coroutine: runs full PPT generation and stores result in _ppt_jobs."""
    store = get_store()
    start_time = time.time()
    try:
        await _do_generate_ppt(job_id, body, auth_header_token, store, start_time)
    except Exception as exc:
        logger.error("PPT background job %s failed: %s", job_id, exc, exc_info=True)
        _ppt_jobs[job_id] = {"status": "error", "error": str(exc)}


async def _do_generate_ppt(job_id: str, body: PPTRequest, auth_header_token: str, store, start_time: float):
    """Core PPT generation logic (extracted so it can run as background task)."""
    # NOTE: job_id is reused as the UUID for the output filename

    # ── 1. Gather context from documents ──
    _loop = asyncio.get_event_loop()
    context_chunks = []
    if body.doc_ids:
        for did in body.doc_ids:
            context_chunks.extend(store.get_chunks_by_doc(did))
    else:
        query_emb = await _loop.run_in_executor(None, embedder.embed_query, body.topic)
        candidates = await _loop.run_in_executor(
            None, store.hybrid_search, body.topic, query_emb, 80
        )
        context_chunks = await _loop.run_in_executor(
            None, rerank, body.topic, candidates, 40
        )

    # llama-server runs 5 parallel slots: TOTAL 16640 tokens / 5 = 3328 tokens per request.
    # Budget: 3328 - 500 (prompt template) - 1500 (output) = ~1300 tokens (~5000 chars) for context.
    # We keep a FULL context buffer (~30000 chars) used in multi-batch mode for richer output;
    # the per-call slice is small enough to stay safely under 3328 tokens.
    context_full_text = "\n\n".join(c["text"] for c in context_chunks[:40])
    if len(context_full_text) > 30000:
        context_full_text = context_full_text[:30000] + "\n[Content truncated]"
    # Single-pass context (used by fallback paths) — safely under 3328-token window.
    context_text = context_full_text[:4500]
    if len(context_full_text) > 4500:
        context_text = context_text + "\n[Content truncated]"

    # ── 2. Extract images from uploaded documents ──
    extracted_images = []
    try:
        from api.generators.image_extractor import get_best_images_for_topic
        doc_ids_for_images = body.doc_ids if body.doc_ids else list({
            c.get("metadata", {}).get("doc_id", "") for c in context_chunks if c.get("metadata", {}).get("doc_id")
        })
        if doc_ids_for_images:
            extracted_images = get_best_images_for_topic(doc_ids_for_images, max_total=5)
            logger.info("Extracted %d images from documents for PPT.", len(extracted_images))
    except Exception as e:
        logger.debug("Image extraction skipped: %s", e)

    has_images_hint = ""
    if extracted_images:
        has_images_hint = f"""
NOTE: {len(extracted_images)} images/diagrams were found in the source documents.
Include {min(len(extracted_images), 3)} slides with "layout": "image" to display them.
Place image slides near relevant content sections."""

    # ── 2.5 Data pre-pass skipped: context is already capped at 4000 chars to fit token window.
    # Adding a pre-pass LLM call consumes extra tokens and causes Gemma echo artifacts.
    extracted_data_hint = ""

    # ── 3. Build LLM prompt ──
    if body.revision_prompt and body.previous_slides_json:
        prompt = f"""You previously generated a PowerPoint about: {body.topic}

Current slides (JSON):
{body.previous_slides_json}

User revision request (this becomes version v{body.version}):
{body.revision_prompt}

Additional context:
{context_text}
{extracted_data_hint}
{has_images_hint}

Return ONLY an updated JSON array with all layout fields preserved and changes applied."""
    else:
        # Plan F: Explicitly prioritize doc_ids context over general knowledge
        doc_priority_note = ""
        if body.doc_ids:
            doc_priority_note = "\nUse ONLY the document content below. No general knowledge.\n"
        prompt = f"""TOPIC: {body.topic}
{doc_priority_note}
DOCUMENT CONTENT:
{context_text}
{has_images_hint}

{_PPT_SLIDE_LAYOUTS}

Rules:
- Exactly {body.num_slides} slides total
- First slide: layout=title (title+subtitle)
- Last slide: layout=thank_you
- Include ONE diagram slide (layout=diagram) with diagram_data
- Include section_header between major sections
- Every slide must have "layout" and "title"
{f'- Style: {body.style_notes}' if body.style_notes else ''}

Return ONLY a valid JSON array of {body.num_slides} slide objects:"""

    # ── 4. Generate slides with retry ──
    # Per-request budget: 3328 tokens. Input (prompt+context) ~1300 tokens.
    # Remaining for output ~2000 tokens. Each slide JSON ~250 tokens.
    # Safe cap: 7 slides single-pass (1400-1750 tokens output).
    # Multi-batch mode (used FIRST below) generates the full body.num_slides via
    # 3 LLM calls of 3-4 slides each — each call uses fresh context slice.
    _safe_slides = min(7, body.num_slides)
    _reduced_slides = min(5, body.num_slides)

    _diagram_hint = (
        'Include ONE slide with layout "diagram" containing a diagram_data object with '
        '"type": "flowchart", "nodes": [{"id":"A","label":"...","shape":"rounded_rect"},...], '
        '"edges": [{"from":"A","to":"B"},...]. This is MANDATORY.'
    )

    # Rewrite main prompt to use safe slide count instead of body.num_slides
    # so the JSON output fits within token budget.
    _prompt_safe = prompt.replace(
        f"Exactly {body.num_slides} slides total", f"Exactly {_safe_slides} slides total"
    ).replace(
        f"Return ONLY a valid JSON array of {body.num_slides} slide objects",
        f"Return ONLY a valid JSON array of {_safe_slides} slide objects"
    )

    _attempt_configs = [
        # (temperature, prompt, num_slides, max_tokens)
        # 3328-token window: input ~1300 tokens, output budget ~2000 tokens.
        # Each slide ~250 tokens; 7 slides = ~1750 tokens. max_tokens=2048 safe.
        (0.1, _prompt_safe, _safe_slides, 2048),
        (
            0.0,
            f"""Create exactly {_safe_slides} slides about: {body.topic}

Context:
{context_text[:3500]}

Return ONLY a valid JSON array. Rules:
- Slide 1: layout="title", must have "title" and "subtitle"
- Last slide: layout="thank_you"
- Other slides: layout="bullets" or "section_header" or "two_column" or "diagram"
- Every slide must have "title" key
- {_diagram_hint}

Return ONLY the JSON array, no markdown, no explanation:""",
            _safe_slides,
            2048,
        ),
        (
            0.0,
            f"""Create {_reduced_slides} slides about: {body.topic}

Context (use ONLY this):
{context_text[:2500]}

Return ONLY a JSON array like:
[
  {{"layout":"title","title":"...","subtitle":"..."}},
  {{"layout":"diagram","title":"...","diagram_data":{{"type":"flowchart","nodes":[{{"id":"A","label":"Step 1","shape":"rounded_rect"}},{{"id":"B","label":"Step 2","shape":"rect"}}],"edges":[{{"from":"A","to":"B"}}]}}}},
  {{"layout":"bullets","title":"Key Points","bullets":["Point 1","Point 2","Point 3"]}},
  {{"layout":"thank_you","title":"Thank You","subtitle":"Indian Coast Guard"}}
]
Use real content from the context above. Return ONLY the JSON array:""",
            _reduced_slides,
            1024,
        ),
    ]

    slides_data = None

    # ═══════════════════════════════════════════════════════════
    # MULTI-BATCH GENERATION (Primary path for RICH content)
    # 3 LLM calls × 3-4 slides each = 10 slides with detailed bullets, tables, diagrams.
    # Each call uses ~1300 input + ~1500 output tokens — safely under 3328 limit.
    # ═══════════════════════════════════════════════════════════
    if not body.revision_prompt and body.num_slides >= 6:
        try:
            slides_data = await _generate_slides_multibatch(
                body, context_full_text, source_filenames=[
                    c.get("metadata", {}).get("filename", "") for c in context_chunks
                ]
            )
            if slides_data and len(slides_data) >= 3:
                logger.info("PPT multi-batch generation: produced %d rich slides", len(slides_data))
            else:
                slides_data = None
        except Exception as e:
            logger.warning("PPT multi-batch generation failed (%s) — falling back to single-pass", e)
            slides_data = None

    if slides_data:
        # Skip the single-pass attempt loop since multi-batch succeeded
        _attempt_configs = []

    for attempt, (temp, prompt_text, n_slides, max_tok) in enumerate(_attempt_configs, 1):
        raw = ""
        try:
            raw = await asyncio.to_thread(
                llm_engine.generate, [{"role": "user", "content": prompt_text}],
                max_tokens=max_tok, temperature=temp, raw=True,
            )

            # If LLM was busy and returned empty/tiny response, wait and retry once
            if len((raw or "").strip()) < 20:
                logger.warning("PPT attempt %d: LLM returned empty/tiny response (%d chars), waiting 3s", attempt, len(raw))
                await asyncio.sleep(3)
                raw = await asyncio.to_thread(
                    llm_engine.generate, [{"role": "user", "content": prompt_text}],
                    max_tokens=max_tok, temperature=temp, raw=True,
                )

            cleaned = _clean_json(raw)
            slides_data = json.loads(cleaned)
            if isinstance(slides_data, list) and len(slides_data) > 0:
                logger.info("PPT JSON parsed successfully on attempt %d (%d slides)", attempt, len(slides_data))
                break
            raise ValueError("Empty or non-array slides")
        except (json.JSONDecodeError, ValueError) as e:
            # Try truncation recovery before giving up on this attempt
            if raw and '[' in raw:
                recovered = _repair_truncated_array(_clean_json(raw) or raw)
                if recovered:
                    try:
                        slides_data = json.loads(recovered)
                        if isinstance(slides_data, list) and len(slides_data) > 0:
                            logger.info("PPT JSON recovered from truncation on attempt %d (%d slides)", attempt, len(slides_data))
                            break
                    except json.JSONDecodeError:
                        pass
            logger.warning("PPT attempt %d failed: %s\nRaw (first 400 chars): %s", attempt, e, raw[:400])
            slides_data = None

    if not slides_data:
        logger.warning("PPT generation failed after all LLM retries. Using fallback slide builder.")
        slides_data = _build_fallback_slides(body.topic, context_text, body.num_slides)

    # ── 5. Flatten nested lists, validate layout fields and enforce exact slide count ──
    def _flatten_slides(data):
        result = []
        for item in data:
            if isinstance(item, list):
                result.extend(_flatten_slides(item))
            elif isinstance(item, dict):
                result.append(item)
        return result
    slides_data = _flatten_slides(slides_data)

    valid_layouts = {"title", "section_header", "bullets", "two_column", "table", "diagram", "chart", "image", "sources", "thank_you"}
    for sd in slides_data:
        layout = sd.get("layout", "bullets")
        if layout not in valid_layouts:
            sd["layout"] = "bullets"
        if "bullets" not in sd:
            sd["bullets"] = []

    # Guarantee at least one diagram slide if none produced by LLM
    has_diagram = any(sd.get("layout") == "diagram" for sd in slides_data)
    if not has_diagram and len(slides_data) >= 3:
        # Build a sensible system-architecture diagram from doc/topic context
        topic_words = [w for w in body.topic.replace('-', ' ').replace('_', ' ').split() if len(w) > 2]
        diagram_title = f"{body.topic} — System Overview" if body.topic else "System Architecture"
        auto_diagram = {
            "layout": "diagram",
            "title": diagram_title,
            "diagram_data": {
                "type": "flowchart",
                "nodes": [
                    {"id": "A", "label": "Document Ingestion", "shape": "rounded_rect"},
                    {"id": "B", "label": "Vector Store / BM25", "shape": "rect"},
                    {"id": "C", "label": "RAG Pipeline", "shape": "rounded_rect"},
                    {"id": "D", "label": "LLM Generation", "shape": "rect"},
                    {"id": "E", "label": "Secure Output", "shape": "rounded_rect"},
                ],
                "edges": [
                    {"from": "A", "to": "B", "label": "index"},
                    {"from": "B", "to": "C", "label": "retrieve"},
                    {"from": "C", "to": "D", "label": "prompt"},
                    {"from": "D", "to": "E", "label": "stream"},
                ],
            },
            "notes": "Auto-generated system architecture diagram",
        }
        # Insert after the second slide (after title/section_header)
        insert_pos = min(2, len(slides_data) - 1)
        slides_data.insert(insert_pos, auto_diagram)
        logger.info("PPT: auto-injected diagram slide at position %d (LLM produced none)", insert_pos)

    # Enforce exact slide count
    target = body.num_slides
    if len(slides_data) < target:
        # Pad with content-bearing bullet slides from context, with meaningful titles
        _pad_sentences = [s.strip() for s in re.split(r'[\n]+|(?<=\.)\s+', context_text) if len(s.strip()) > 40]
        _pad_idx = 0
        _slide_num = len(slides_data) + 1
        while len(slides_data) < target:
            _chunk = _pad_sentences[_pad_idx:_pad_idx+4] or ["Refer to source document for additional technical details."]
            _pad_idx = (_pad_idx + 4) % max(1, len(_pad_sentences))
            # Extract a meaningful title from first sentence or use section-based title
            _title_source = _chunk[0][:60] if _chunk else "Additional Information"
            _section_title = _extract_section_title(_title_source) or f"Section {_slide_num - 1}"
            slides_data.insert(max(1, len(slides_data) - 1), {
                "layout": "bullets",
                "title": _section_title,
                "bullets": [s[:180] for s in _chunk]
            })
            _slide_num += 1
    elif len(slides_data) > target:
        # Keep first (title) and last (thank_you); trim from middle
        first = slides_data[0]
        last = slides_data[-1]
        middle = slides_data[1:-1]
        keep = middle[:max(0, target - 2)]
        slides_data = [first] + keep + [last]

    # Auto-inject sources slide if documents were used and count allows
    source_names = []
    if body.doc_ids:
        source_names = [store.get_chunks_by_doc(did)[0]["metadata"].get("filename", did) for did in body.doc_ids if store.get_chunks_by_doc(did)]
    else:
        source_names = list({c.get("metadata", {}).get("filename", "") for c in context_chunks if c.get("metadata", {}).get("filename")})
    source_names = [s for s in source_names if s]
    if source_names and len(slides_data) >= 3:
        # Replace second-to-last slide (before thank_you) with sources
        slides_data[-2] = {"layout": "sources", "title": "Sources & References", "sources": source_names[:20]}

    # ═══════════════════════════════════════════════════════════════════════
    # MODULE 5 & 10: Check for superseded documents and add genealogy info
    # ═══════════════════════════════════════════════════════════════════════
    superseded_warnings = []
    genealogy_data = []
    
    try:
        # Get doc_ids from context
        doc_ids_for_genealogy = body.doc_ids if body.doc_ids else list({
            c.get("metadata", {}).get("doc_id", "") for c in context_chunks if c.get("metadata", {}).get("doc_id")
        })
        
        if doc_ids_for_genealogy and should_include_genealogy(doc_ids_for_genealogy):
            # Check superseded status
            superseded_status = await check_superseded_status(doc_ids_for_genealogy, auth_header_token)
            
            if superseded_status:
                # Build warnings list
                for doc_id, info in superseded_status.items():
                    superseded_warnings.append({
                        "old_doc": info.get("superseded_by_name", "Unknown"),
                        "new_doc": info.get("superseded_by_name", "Newer Version"),
                        "date": info.get("date", ""),
                    })
                
                logger.info("PPT: Found %d superseded documents, adding warning slide", len(superseded_warnings))
            
            # Get genealogy for each doc
            for doc_id in doc_ids_for_genealogy[:5]:  # Limit to first 5 docs
                lineage = await get_document_lineage(doc_id, auth_header_token)
                if lineage:
                    genealogy_data.append(lineage)
            
            if genealogy_data:
                logger.info("PPT: Retrieved genealogy for %d documents", len(genealogy_data))
    except Exception as e:
        logger.warning("PPT: Failed to fetch genealogy info: %s", e)
    
    # Insert superseded warning slide as slide 2 (after title) if warnings exist
    if superseded_warnings and len(slides_data) >= 2:
        warning_slide = {
            "layout": "superseded_warning",
            "title": "Document Status Warning",
            "warnings": superseded_warnings,
        }
        slides_data.insert(1, warning_slide)  # Insert after title slide
        logger.info("PPT: Inserted superseded warning slide at position 1")
    
    # Insert genealogy slide before sources slide (if genealogy data exists)
    if genealogy_data and len(slides_data) >= 3:
        genealogy_slide = {
            "layout": "genealogy",
            "title": "Document Genealogy & Provenance",
            "genealogy_data": genealogy_data,
        }
        # Insert before the last two slides (sources + thank_you)
        insert_pos = len(slides_data) - 2 if len(slides_data) >= 4 else len(slides_data) - 1
        slides_data.insert(insert_pos, genealogy_slide)
        logger.info("PPT: Inserted genealogy slide at position %d", insert_pos)

    # ── Post-process: clean slide content ──
    # Remove footer text, source trace bullets, and other artifacts
    _FOOTER_PATTERNS = [
        r'AI-Generated Draft\s*\|\s*Indian Coast Guard\s*\|\s*AGRA System\s*\|\s*Confidential',
        r'\[Document:.*?\| Page \d+\]',
        r'Generated by AGRA.*?System',
    ]
    for slide in slides_data:
        # Clean bullets
        if slide.get("bullets"):
            cleaned_bullets = []
            for bullet in slide["bullets"]:
                # Skip source trace bullets
                if re.search(r'\[Document:.*?\|', bullet):
                    continue
                # Remove footer patterns from bullet text
                for pattern in _FOOTER_PATTERNS:
                    bullet = re.sub(pattern, '', bullet, flags=re.IGNORECASE)
                bullet = bullet.strip()
                if bullet and len(bullet) > 5:
                    cleaned_bullets.append(bullet)
            slide["bullets"] = cleaned_bullets
        # Clean two_column content
        if slide.get("layout") == "two_column":
            for key in ["left_column", "right_column", "left_items", "right_items"]:
                if slide.get(key):
                    cleaned = []
                    for item in slide[key]:
                        if re.search(r'\[Document:.*?\|', item):
                            continue
                        for pattern in _FOOTER_PATTERNS:
                            item = re.sub(pattern, '', item, flags=re.IGNORECASE)
                        item = item.strip()
                        if item and len(item) > 5:
                            cleaned.append(item)
                    slide[key] = cleaned

    # ── 6. Build PPTX ──
    # job_id is passed in as parameter (matches the polling key in _ppt_jobs)
    version_label = f"_v{body.version}" if body.version > 1 else ""
    # When doc_ids are provided and topic is a generic fallback or still looks like
    # an unparsed user query, derive a meaningful presentation title from the first
    # document's filename (strip extension).
    presentation_title = body.topic
    if body.doc_ids and context_chunks:
        first_filename = context_chunks[0].get("metadata", {}).get("filename", "")
        if first_filename and (
            not body.topic
            or body.topic.lower() in ("document overview", "")
            or re.match(r'^(can|could|would|should|how|what|why|when|where|do|does|did|is|are|will)', body.topic, re.IGNORECASE)
            or len(body.topic.split()) > 8
        ):
            stem = Path(first_filename).stem.replace('_', ' ').replace('-', ' ')
            presentation_title = stem[:80] or body.topic
        safe_topic = re.sub(r'[^\w\s-]', '', first_filename or body.topic)[:30].replace(' ', '_')
    else:
        safe_topic = re.sub(r'[^\w\s-]', '', body.topic)[:30].replace(' ', '_')
    safe_topic = safe_topic or "Presentation"
    version_label = f"_v{body.version}" if body.version > 1 else "_v1"
    output_filename = f"{safe_topic}{version_label}.pptx"
    output_path = _OUTPUTS_DIR / output_filename

    # Patch the title slide with the resolved presentation_title
    if slides_data and slides_data[0].get("layout") == "title":
        slides_data[0]["title"] = presentation_title

    # ── 6.5 ICG Master Template Integration ──
    assets_dir = Path(__file__).resolve().parent.parent.parent / "assets"
    template_path = str(assets_dir / "icg_master.pptx") if (assets_dir / "icg_master.pptx").exists() else None

    build_pptx(
        slides_data,
        str(output_path),
        title=presentation_title,
        template_path=template_path,
        extracted_images=extracted_images,
    )

    logger.info("Generated PPT v%d: %s (%d slides, %d doc images)", body.version, output_path.name, len(slides_data), len(extracted_images))

    elapsed_ms = (time.time() - start_time) * 1000
    log_usage(
        action_type="ppt",
        module="generate",
        token=auth_header_token,
        response_time_ms=elapsed_ms,
    )

    # Store result in job store so status endpoint can serve it
    _ppt_jobs[job_id] = {
        "status": "done",
        "filename": output_filename,
        "pretty_filename": output_filename,
        "slides_json": json.dumps(slides_data),
        "version": body.version,
    }


@router.post("/generate/ppt")
async def generate_ppt(
    body: PPTRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Start async PPT generation. Returns {job_id} immediately (no timeout risk).
    Frontend polls GET /generate/ppt/status/{job_id} every 5s for completion.
    """
    auth_header_token = ""
    if request:
        auth_h = request.headers.get("authorization", "")
        auth_header_token = auth_h.replace("Bearer ", "") if auth_h else ""

    # ── RBAC Check: Verify user can access all requested documents ──
    from api.utils.auth_check import can_access_document
    if body.doc_ids:
        store = get_store()
        for did in body.doc_ids:
            doc_meta = store.get_document_metadata(did)
            if doc_meta:
                doc_clearance = doc_meta.get("clearance_level", 1)
                if not can_access_document(user, doc_clearance):
                    raise HTTPException(
                        status_code=403,
                        detail=f"Access denied: Document '{did}' requires clearance level {doc_clearance}"
                    )

    job_id = str(uuid.uuid4())
    _ppt_jobs[job_id] = {"status": "pending"}

    asyncio.create_task(_run_ppt_job(job_id, body, auth_header_token))

    return JSONResponse({"job_id": job_id, "status": "pending"})


# ═══════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

class SummaryRequest(BaseModel):
    doc_ids: List[str] = Field(default_factory=list)
    doc_id: Optional[str] = None  # Legacy support
    summary_type: str = Field(default="executive", pattern="^(executive|technical)$")
    detail_level: str = Field(default="detailed", pattern="^(brief|detailed)$")


@router.post("/generate/summary")
async def generate_summary(
    body: SummaryRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Generate an executive or technical summary of a document.
    Returns SSE stream of the summary + a .docx download link.
    """
    summary_start = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""
    store = get_store()
    target_doc_ids = body.doc_ids if body.doc_ids else ([body.doc_id] if body.doc_id else [])
    if not target_doc_ids:
        raise HTTPException(status_code=400, detail="Must provide at least one doc_id.")

    # ── RBAC Check: Verify user can access all requested documents ──
    from api.utils.auth_check import can_access_document, get_user_clearance, is_superadmin
    for did in target_doc_ids:
        doc_meta = store.get_document_metadata(did)
        if doc_meta:
            doc_clearance = doc_meta.get("clearance_level", 1)
            if not can_access_document(user, doc_clearance):
                raise HTTPException(
                    status_code=403,
                    detail=f"Access denied: Document '{did}' requires clearance level {doc_clearance}"
                )

    chunks = []
    filenames = []
    doc_chunks_map = {}  # Map doc_id to its chunks
    
    for did in target_doc_ids:
        doc_chunks = store.get_chunks_by_doc(did)
        if doc_chunks:
            chunks.extend(doc_chunks)
            doc_chunks_map[did] = doc_chunks
            if doc_chunks[0]["metadata"].get("filename") not in filenames:
                filenames.append(doc_chunks[0]["metadata"].get("filename", "document"))

    if not chunks:
        raise HTTPException(status_code=404, detail="Documents not found in knowledge base.")

    # ═══════════════════════════════════════════════════════════════════════
    # MODULE 5 & 10: Check for superseded documents and get genealogy
    # ═══════════════════════════════════════════════════════════════════════
    superseded_warning_text = ""
    genealogy_data = []
    
    try:
        if should_include_genealogy(target_doc_ids):
            # Check superseded status
            superseded_status = await check_superseded_status(target_doc_ids, auth_tok)
            if superseded_status:
                superseded_warning_text = format_superseded_warning(superseded_status)
                logger.info("Summary: Found %d superseded documents", len(superseded_status))
            
            # Get genealogy for each doc
            for doc_id in target_doc_ids[:5]:
                lineage = await get_document_lineage(doc_id, auth_tok)
                if lineage:
                    genealogy_data.append(lineage)
    except Exception as e:
        logger.warning("Summary: Failed to fetch genealogy info: %s", e)

    # Build hierarchical context (Module 5: Multi-document hierarchical summarization)
    # Level 1: Per-document summaries
    doc_contexts = []
    for doc_id, doc_chunks in doc_chunks_map.items():
        doc_text = "\n\n".join(c["text"] for c in doc_chunks[:5])
        doc_filename = doc_chunks[0]["metadata"].get("filename", doc_id)
        doc_contexts.append(f"--- {doc_filename} ---\n{doc_text[:2000]}")
    
    # Combine with full text for comprehensive summary
    full_text = "\n\n".join(c["text"] for c in chunks[:10])
    if len(full_text) > 8000:
        full_text = full_text[:8000] + "\n[Content truncated for summary generation]"

    filename_label = ", ".join(filenames) if len(filenames) <= 3 else f"{len(filenames)} Documents"
    
    # Generate safe topic for filename
    first_filename = filenames[0] if filenames else "document"
    safe_topic = re.sub(r'[^\w\s-]', '', Path(first_filename).stem)[:30].replace(' ', '_') or "document"

    type_label = "Executive Summary" if body.summary_type == "executive" else "Technical Summary"
    
    # Include superseded warning in prompt if applicable
    superseded_note = ""
    if superseded_warning_text:
        superseded_note = f"\n\nIMPORTANT - DOCUMENT STATUS:\n{superseded_warning_text}\n\n"
    
    detail_instruction = ""
    if body.detail_level == "brief":
        detail_instruction = (
            "Keep the summary CONCISE — maximum 3 paragraphs per document.\n"
            "Focus ONLY on the most critical findings, key numbers, and bottom-line conclusions.\n"
            "Omit background explanation and peripheral details.\n"
            "Target length: 300-500 words total."
        )
    else:
        detail_instruction = (
            "Provide a COMPREHENSIVE and THOROUGH summary.\n"
            "Include relevant details, data points, exceptions, and nuances.\n"
            "Cover each document's scope, methodology, findings, and implications.\n"
            "Target length: 1000-2000 words total."
        )
    
    prompt = f"""Generate a {type_label} of the following document(s). {detail_instruction}
{superseded_note}
DOCUMENTS: {filename_label}

CONTENT:
{full_text}

FORMAT:
1. **Overview** — 2-3 sentences describing the document's purpose and scope
2. **Document Status** — Note any superseded documents and their replacements
3. **Key Points by Document** — Bullet list of most important findings per document
4. **Cross-Document Analysis** — Compare and synthesize information across sources
5. **Conclusions & Recommendations** — Actionable takeaways with citations

CITE using format: [Document Name, p.X] or [Document Name, Section Y]."""

    messages = [
        {"role": "user", "content": f"You are a senior analyst creating professional document summaries for Indian Coast Guard leadership.\n\n{prompt}"},
    ]

    # Stream the summary and also collect for DOCX
    job_id = str(uuid.uuid4())
    collected_text: list = []

    async def event_stream():
        token_queue = asyncio.Queue()

        def _run_llm():
            try:
                for tok in llm_engine.stream_generate(messages, max_tokens=2048):
                    token_queue.put_nowait(tok)
                token_queue.put_nowait(None)
            except Exception as e:
                logger.error("Summary LLM error: %s", e)
                token_queue.put_nowait(None)

        llm_thread = asyncio.get_event_loop().run_in_executor(None, _run_llm)

        while True:
            tok = await token_queue.get()
            if tok is None:
                break
            collected_text.append(tok)
            yield f"data: {json.dumps({'token': tok})}\n\n"

        # Build DOCX
        summary_filename = f"{safe_topic}_summary_v1.docx"
        docx_path = _OUTPUTS_DIR / summary_filename
        doc = DocxDocument()
        doc.add_heading(f"{type_label}: {filename_label}", level=1)
        
        # Add superseded warning if applicable (Module 10)
        if superseded_warning_text:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run(superseded_warning_text)
            warning_run.font.color.rgb = None  # Use default (usually red via markdown)
            warning_run.bold = True
            doc.add_paragraph()  # Spacing
        
        doc.add_paragraph("".join(collected_text))
        
        # Add Document Genealogy section (Module 5 & 10)
        if genealogy_data:
            doc.add_page_break()
            doc.add_heading("Document Genealogy & Provenance", level=2)
            
            # Create table for genealogy info
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Document'
            hdr_cells[1].text = 'Version'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Relationships'
            
            # Data rows
            for lineage in genealogy_data:
                row_cells = table.add_row().cells
                row_cells[0].text = lineage.get('filename', 'Unknown')
                row_cells[1].text = f"v{lineage.get('version', '?')}"
                row_cells[2].text = lineage.get('status', 'unknown')
                
                # Build relationship text
                rel_parts = []
                if lineage.get('superseded_by_name'):
                    rel_parts.append(f"Superseded by: {lineage['superseded_by_name']}")
                if lineage.get('supersedes'):
                    supersedes_list = ", ".join(s.get('filename', '?') for s in lineage['supersedes'])
                    rel_parts.append(f"Supersedes: {supersedes_list}")
                
                row_cells[3].text = "; ".join(rel_parts) if rel_parts else "None"
            
            doc.add_paragraph()
        
        # Add Watermark (FR-GEN-006)
        try:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "AI-Generated Draft - ICG AGRA"
        except Exception as e:
            logger.warning(f"Could not add watermark: {e}")

        doc.save(str(docx_path))

        # Build PDF export
        pdf_filename = f"{safe_topic}_summary_v1.pdf"
        pdf_path = _OUTPUTS_DIR / pdf_filename
        try:
            full_summary_text = "".join(collected_text)
            generate_summary_pdf(
                title=f"{type_label}: {filename_label}",
                content_text=full_summary_text,
                output_path=pdf_path,
                detail_level=body.detail_level,
            )
        except Exception as e:
            logger.warning("Summary PDF generation failed: %s", e)
            pdf_filename = None

        result = {'done': True, 'download_url': f'/api/agent/download/{summary_filename}'}
        if pdf_filename:
            result['pdf_download_url'] = f'/api/agent/download/{pdf_filename}'
        yield f"data: {json.dumps(result)}\n\n"

        # Log usage
        elapsed_ms = (time.time() - summary_start) * 1000
        log_usage(action_type="summary", module="generate", token=auth_tok, response_time_ms=elapsed_ms)

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ═══════════════════════════════════════════════════════════════
#  KNOWLEDGE QUIZ
# ═══════════════════════════════════════════════════════════════

class QuizRequest(BaseModel):
    doc_id: str
    num_mcq: int = Field(default=5, ge=0, le=20)
    num_true_false: int = Field(default=3, ge=0, le=10)
    num_short_answer: int = Field(default=2, ge=0, le=10)
    difficulty: str = Field(default="medium", description="easy, medium, hard")
    scope: str = Field(default="comprehensive", description="concepts, details, comprehensive")


@router.post("/generate/quiz")
async def generate_quiz(
    body: QuizRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Generate a knowledge quiz from a document.
    Returns JSON quiz data + a .docx download link.
    """
    quiz_start = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""

    # ── RBAC Check: Verify user can access the requested document ──
    from api.utils.auth_check import can_access_document
    store = get_store()
    doc_meta = store.get_document_metadata(body.doc_id)
    if doc_meta:
        doc_clearance = doc_meta.get("clearance_level", 1)
        if not can_access_document(user, doc_clearance):
            raise HTTPException(
                status_code=403,
                detail=f"Access denied: Document '{body.doc_id}' requires clearance level {doc_clearance}"
            )

    chunks = store.get_chunks_by_doc(body.doc_id)

    if not chunks:
        raise HTTPException(status_code=404, detail=f"Document {body.doc_id} not found in knowledge base.")

    # Plan F: Tighten quiz context to fit 3328-token window (~3000 chars ≈ 750 tokens)
    content = "\n\n".join(c["text"] for c in chunks[:8])
    if len(content) > 3000:
        content = content[:3000] + "\n[Content truncated for quiz generation]"
    filename = chunks[0]["metadata"].get("filename", "document")

    # Build question section lines based on counts > 0
    _sections = []
    if body.num_mcq > 0:
        _sections.append(f"- {body.num_mcq} Multiple Choice Questions (MCQ) with 4 options each (A, B, C, D) and the correct answer key")
    if body.num_true_false > 0:
        _sections.append(f"- {body.num_true_false} True/False Questions with the correct answer (True or False) and a brief explanation")
    if body.num_short_answer > 0:
        _sections.append(f"- {body.num_short_answer} Short Answer Questions with model answers")
    _sections_str = "\n".join(_sections)

    prompt = f"""Generate a knowledge assessment quiz from this document.

DOCUMENT: {filename}

CONTENT:
{content}

Generate EXACTLY:
{_sections_str}

DIFFICULTY LEVEL: {body.difficulty.upper()}
SCOPE FOCUS: {body.scope.upper()}

Return ONLY valid JSON in this exact format:
{{
  "title": "Quiz: [document topic]",
  "mcq": [
    {{
      "question": "...",
      "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
      "correct": "A",
      "explanation": "..."
    }}
  ],
  "true_false": [
    {{
      "question": "...",
      "answer": true,
      "explanation": "..."
    }}
  ],
  "short_answer": [
    {{
      "question": "...",
      "model_answer": "..."
    }}
  ]
}}"""

    messages = [
        {"role": "user", "content": f"You are an expert quiz designer. Return only valid JSON.\n\n{prompt}"},
    ]

    raw = await asyncio.to_thread(
        llm_engine.generate, messages,
        max_tokens=4096, temperature=0.3, raw=True,
    )

    quiz_data: Optional[dict] = None
    last_error: Optional[Exception] = None

    # Pass 1: clean + parse with depth-counting matcher
    try:
        cleaned = _clean_json(raw)
        if not cleaned:
            raise ValueError("Empty response from LLM")
        quiz_data = json.loads(cleaned)
    except (json.JSONDecodeError, ValueError) as e:
        last_error = e
        logger.warning("Quiz JSON parse pass 1 failed: %s", e)

    # Pass 1b: Truncation recovery — try to salvage partial JSON
    if quiz_data is None and cleaned:
        try:
            # Find last complete MCQ/TF/SA object and close the structure
            for marker in ['"explanation"', '"model_answer"', '"answer"']:
                idx = cleaned.rfind(marker)
                if idx > 0:
                    # Find the closing } of this object
                    close_brace = cleaned.find('}', idx)
                    if close_brace > 0:
                        candidate = cleaned[:close_brace + 1]
                        # Close any open arrays/objects
                        open_brackets = sum(
                            1 if c == '[' else -1 if c == ']' else 0
                            for c in candidate
                        )
                        open_braces = sum(
                            1 if c == '{' else -1 if c == '}' else 0
                            for c in candidate
                        )
                        candidate += ']' * max(0, open_brackets)
                        candidate += '}' * max(0, open_braces)
                        quiz_data = json.loads(candidate)
                        logger.info("Quiz JSON recovered via truncation recovery (marker: %s)", marker)
                        break
        except (json.JSONDecodeError, ValueError):
            pass

    # Pass 2: LLM repair if pass 1 failed
    if quiz_data is None:
        logger.info("Attempting LLM-based JSON repair for quiz output...")
        repaired = await asyncio.to_thread(
            _repair_json_with_llm, raw, str(last_error or "unknown")
        )
        if repaired is not None:
            quiz_data = repaired
            logger.info("Quiz JSON repaired successfully via LLM pass.")

    # Pass 3: Final synthetic fallback to avoid 500 error
    if quiz_data is None or not isinstance(quiz_data, dict):
        logger.error(
            "Quiz JSON unrecoverable. Returning minimal fallback. Raw (first 800 chars): %s",
            raw[:800],
        )
        quiz_data = {
            "title": f"Quiz: {filename}",
            "mcq": [],
            "true_false": [],
            "short_answer": [],
            "_fallback": True,
            "_error": "The model failed to produce structured questions. Please try again.",
        }

    # Validate minimal structure
    if "mcq" not in quiz_data:
        quiz_data["mcq"] = []
    if "true_false" not in quiz_data:
        quiz_data["true_false"] = []
    if "short_answer" not in quiz_data:
        quiz_data["short_answer"] = []

    # ═══════════════════════════════════════════════════════════════════════
    # MODULE 10: Check for superseded document status
    # ═══════════════════════════════════════════════════════════════════════
    superseded_warning_text = ""
    try:
        if should_include_genealogy([body.doc_id]):
            superseded_status = await check_superseded_status([body.doc_id], auth_tok)
            if superseded_status:
                superseded_warning_text = format_superseded_warning(superseded_status)
                logger.info("Quiz: Source document is superseded, adding warning")
    except Exception as e:
        logger.warning("Quiz: Failed to check superseded status: %s", e)

    # Build DOCX
    job_id = str(uuid.uuid4())
    quiz_topic = re.sub(r'[^\w\s-]', '', Path(filename).stem)[:30].replace(' ', '_') or "quiz"
    quiz_filename = f"{quiz_topic}_quiz_v1.docx"
    docx_path = _OUTPUTS_DIR / quiz_filename
    doc = DocxDocument()
    doc.add_heading(quiz_data.get("title", f"Quiz: {filename}"), level=1)
    
    # Add superseded warning banner if applicable (Module 10)
    if superseded_warning_text:
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run("⚠️ " + superseded_warning_text.replace("⚠️ DOCUMENT STATUS WARNING:\n", ""))
        warning_run.bold = True
        warning_run.font.size = None  # Default size
        doc.add_paragraph()

    doc.add_heading("Multiple Choice Questions", level=2)
    for i, q in enumerate(quiz_data.get("mcq", []), 1):
        if not isinstance(q, dict):
            continue
        doc.add_paragraph(f"Q{i}. {q.get('question', '')}")
        options = q.get("options")
        if not isinstance(options, dict):
            options = {}
        for key in ("A", "B", "C", "D"):
            opt = options.get(key, "")
            doc.add_paragraph(f"   {key}) {opt}")
        doc.add_paragraph(f"   ✓ Correct: {q.get('correct', '?')}")
        doc.add_paragraph("")

    if quiz_data.get("true_false"):
        doc.add_heading("True / False Questions", level=2)
        for i, q in enumerate(quiz_data["true_false"], 1):
            if not isinstance(q, dict):
                continue
            ans = "True" if q.get("answer") is True else "False"
            doc.add_paragraph(f"Q{i}. {q.get('question', '')}")
            doc.add_paragraph(f"   ✓ Answer: {ans}")
            if q.get("explanation"):
                doc.add_paragraph(f"   Explanation: {q.get('explanation')}")
            doc.add_paragraph("")

    doc.add_heading("Short Answer Questions", level=2)
    for i, q in enumerate(quiz_data.get("short_answer", []), 1):
        if not isinstance(q, dict):
            continue
        doc.add_paragraph(f"Q{i}. {q.get('question', '')}")
        doc.add_paragraph(f"   Model Answer: {q.get('model_answer', '')}")
        doc.add_paragraph("")

    # Add Watermark (FR-GEN-006)
    try:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "AI-Generated Draft - ICG AGRA"
    except Exception as e:
        logger.warning(f"Could not add watermark: {e}")

    doc.save(str(docx_path))

    # Build PDF export
    quiz_pdf_filename = f"{quiz_topic}_quiz_v1.pdf"
    quiz_pdf_path = _OUTPUTS_DIR / quiz_pdf_filename
    try:
        generate_quiz_pdf(
            quiz_data=quiz_data,
            filename=quiz_filename,
            output_path=quiz_pdf_path,
        )
    except Exception as e:
        logger.warning("Quiz PDF generation failed: %s", e)
        quiz_pdf_filename = None

    elapsed_ms = (time.time() - quiz_start) * 1000
    log_usage(action_type="quiz", module="generate", token=auth_tok, response_time_ms=elapsed_ms)

    result = {
        "quiz": quiz_data,
        "download_url": f"/api/agent/download/{quiz_filename}",
    }
    if quiz_pdf_filename:
        result["pdf_download_url"] = f"/api/agent/download/{quiz_pdf_filename}"
    return result


# ═══════════════════════════════════════════════════════════════
#  DRAFT SOTR GENERATOR (FR-GEN-003)
# ═══════════════════════════════════════════════════════════════

class DraftSOTRRequest(BaseModel):
    doc_id: str
    focus_area: Optional[str] = Field(None, description="Specific area to focus the SOTR on")


@router.post("/generate/sotr")
async def generate_draft_sotr(
    body: DraftSOTRRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Generate a Draft SOTR (Statement of Technical Requirements).
    Returns SSE stream of the SOTR + a .docx download link.
    """
    sotr_start = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""
    
    store = get_store()
    chunks = store.get_chunks_by_doc(body.doc_id)

    if not chunks:
        raise HTTPException(status_code=404, detail=f"Document {body.doc_id} not found in knowledge base.")

    full_text = "\n\n".join(c["text"] for c in chunks)
    # Plan F: Cap SOTR context for 3328-token model
    if len(full_text) > 4000:
        full_text = full_text[:4000] + "\n[Document truncated]"

    filename = chunks[0]["metadata"].get("filename", "document")
    focus_note = f"\\nFocus specifically on: {body.focus_area}" if body.focus_area else ""

    prompt = f"""Generate a Draft Statement of Technical Requirements (SOTR) based on the following document.

DOCUMENT: {filename}
{focus_note}

CONTENT:
{full_text}

FORMAT the SOTR strictly with the following sections:
1. **Introduction** — Purpose and scope
2. **Applicable Documents** — References and standards
3. **General Requirements** — High-level technical needs
4. **Specific Requirements** — Detailed specifications and performance criteria
5. **Quality Assurance & Testing** — Acceptance criteria

Use formal, objective military specification language (e.g., 'The system shall...', 'The contractor must...')."""

    messages = [
        {"role": "system", "content": "You are a technical specifications writer for the Indian Coast Guard."},
        {"role": "user", "content": prompt},
    ]

    job_id = str(uuid.uuid4())
    collected_text = []

    async def event_stream():
        token_queue = asyncio.Queue()

        def _run_llm():
            try:
                for tok in llm_engine.stream_generate(messages, max_tokens=2048):
                    token_queue.put_nowait(tok)
                token_queue.put_nowait(None)
            except Exception as e:
                logger.error("SOTR LLM error: %s", e)
                token_queue.put_nowait(None)

        llm_thread = asyncio.get_event_loop().run_in_executor(None, _run_llm)

        while True:
            tok = await token_queue.get()
            if tok is None:
                break
            collected_text.append(tok)
            yield f"data: {json.dumps({'token': tok})}\n\n"

        sotr_topic = re.sub(r'[^\w\s-]', '', Path(filename).stem)[:30].replace(' ', '_') or "sotr"
        sotr_filename = f"{sotr_topic}_sotr_v1.docx"
        docx_path = _OUTPUTS_DIR / sotr_filename
        doc = DocxDocument()
        doc.add_heading(f"Draft SOTR: {filename}", level=1)
        doc.add_paragraph("".join(collected_text))

        try:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "AI-Generated Draft - ICG AGRA"
        except Exception as e:
            logger.warning(f"Could not add watermark: {e}")

        doc.save(str(docx_path))
        yield f"data: {json.dumps({'done': True, 'download_url': f'/api/agent/download/{sotr_filename}'})}\n\n"

        log_usage(action_type="draft_sotr", module="generate", token=auth_tok, response_time_ms=(time.time() - sotr_start) * 1000)

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ═══════════════════════════════════════════════════════════════
#  TECHNICAL REVIEW COMMENT GENERATOR (FR-GEN-004)
# ═══════════════════════════════════════════════════════════════

class TechReviewRequest(BaseModel):
    doc_id: str
    target_audience: str = Field(default="shipyard", pattern="^(shipyard|internal|management)$")


@router.post("/generate/tech-review")
async def generate_tech_review(
    body: TechReviewRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Generate Technical Review Comments for a document.
    """
    review_start = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""
        
    store = get_store()
    chunks = store.get_chunks_by_doc(body.doc_id)

    if not chunks:
        raise HTTPException(status_code=404, detail="Document not found.")

    full_text = "\n\n".join(c["text"] for c in chunks[:30])
    # Plan F: Cap tech review context for 3328-token model
    if len(full_text) > 4000:
        full_text = full_text[:4000] + "\n[Document truncated]"
    filename = chunks[0]["metadata"].get("filename", "document")

    prompt = f"""Generate Technical Review Comments for the following submission document.
Target Audience: {body.target_audience}

DOCUMENT: {filename}
CONTENT:
{full_text}

FORMAT the review strictly as:
1. **General Observations** — Overall assessment of the submission
2. **Major Deviations / Concerns** — Critical issues identified
3. **Clarifications Required** — Specific points needing more details
4. **Recommendations** — Proposed actions or accept/reject advice

Keep the tone professional, objective, and analytical."""

    messages = [
        {"role": "system", "content": "You are a lead technical reviewer for the Indian Coast Guard."},
        {"role": "user", "content": prompt},
    ]

    job_id = str(uuid.uuid4())
    collected_text = []

    async def event_stream():
        token_queue = asyncio.Queue()

        def _run_llm():
            try:
                for tok in llm_engine.stream_generate(messages, max_tokens=2048):
                    token_queue.put_nowait(tok)
                token_queue.put_nowait(None)
            except Exception as e:
                logger.error("Tech review LLM error: %s", e)
                token_queue.put_nowait(None)

        llm_thread = asyncio.get_event_loop().run_in_executor(None, _run_llm)

        while True:
            tok = await token_queue.get()
            if tok is None:
                break
            collected_text.append(tok)
            yield f"data: {json.dumps({'token': tok})}\n\n"

        review_topic = re.sub(r'[^\w\s-]', '', Path(filename).stem)[:30].replace(' ', '_') or "review"
        review_filename = f"{review_topic}_tech_review_v1.docx"
        docx_path = _OUTPUTS_DIR / review_filename
        doc = DocxDocument()
        doc.add_heading(f"Technical Review: {filename}", level=1)
        doc.add_paragraph("".join(collected_text))

        try:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "AI-Generated Draft - ICG AGRA"
        except Exception as e:
            logger.warning(f"Could not add watermark: {e}")

        doc.save(str(docx_path))
        yield f"data: {json.dumps({'done': True, 'download_url': f'/api/agent/download/{review_filename}'})}\n\n"
        log_usage(action_type="tech_review", module="generate", token=auth_tok, response_time_ms=(time.time() - review_start) * 1000)

    return StreamingResponse(event_stream(), media_type="text/event-stream")
