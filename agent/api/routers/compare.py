"""
AGRA Phase 4 — Router: Cross-Document Reasoning
Multi-document comparative analysis (e.g. Bidder A vs Bidder B vs Standard).

Two endpoints:
  POST /compare/bids        — legacy doc_id-based comparison (flat retrieval)
  POST /compare/bids/branch — branch-isolated comparison using bidder_key /
                              problem_statement metadata (Phase D)
"""

import asyncio
import json
import logging
import re
import time
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from docx import Document as DocxDocument

from api.utils.auth_check import get_current_user
from api.utils.usage_logger import log_usage
from api.rag import llm as llm_engine, embedder
from api.rag.vector_store import get_store

logger = logging.getLogger("agra.compare")

router = APIRouter()

import os as _os
_DATA_DIR = Path(_os.environ.get("AGRA_DATA_DIR", "/workspace/agra_data"))
if not _DATA_DIR.exists():
    _DATA_DIR = Path(__file__).resolve().parent.parent.parent / "agra_data"
_OUTPUTS_DIR = _DATA_DIR / "outputs"
_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)


def _clean_json(raw: str) -> str:
    cleaned = re.sub(r'^```(?:json)?\s*', '', raw.strip(), flags=re.MULTILINE)
    cleaned = re.sub(r'```\s*$', '', cleaned.strip(), flags=re.MULTILINE)
    return cleaned.strip()


class CompareRequest(BaseModel):
    bid_doc_ids: List[str] = Field(..., min_length=2, description="At least two bid documents to compare")
    standard_doc_id: Optional[str] = Field(None, description="Optional standard/SOTR to compare against")
    check_scope: Optional[str] = Field(None, description="Specific area to focus on")


@router.post("/compare/bids")
async def compare_bids(
    body: CompareRequest,
    user: dict = Depends(get_current_user),
    request: Request = None,
):
    """
    Perform a comparative analysis across multiple bid documents.
    If standard_doc_id is provided, evaluates both against the standard.
    Returns SSE stream of comparison matrices.
    """
    compare_start = time.time()
    auth_tok = ""
    if request:
        ah = request.headers.get("authorization", "")
        auth_tok = ah.replace("Bearer ", "") if ah else ""
    store = get_store()

    # Load Standard
    standard_text = ""
    if body.standard_doc_id:
        std_chunks = store.get_chunks_by_doc(body.standard_doc_id)
        if not std_chunks:
            raise HTTPException(status_code=404, detail="Standard document not found.")
        standard_text = "\n\n".join(c["text"] for c in std_chunks[:20])

    # Load Bids
    bids_data = {}
    for bid_id in body.bid_doc_ids:
        chunks = store.get_chunks_by_doc(bid_id)
        if not chunks:
            raise HTTPException(status_code=404, detail=f"Bid document {bid_id} not found.")
        filename = chunks[0]["metadata"].get("filename", f"Bid_{bid_id}")
        bids_data[filename] = "\n\n".join(c["text"] for c in chunks[:20])

    if len(bids_data) < 2:
        raise HTTPException(status_code=400, detail="Must have at least two valid bid documents.")

    bids_context = ""
    for filename, text in bids_data.items():
        bids_context += f"\n--- BID: {filename} ---\n{text[:8000]}\n"

    scope_note = f"\nFocus specifically on: {body.check_scope}" if body.check_scope else ""
    std_note = f"\nSTANDARD / SOTR:\n{standard_text[:8000]}\n" if standard_text else ""

    prompt = f"""You are a senior procurement evaluation officer for the Indian Coast Guard.

TASK: Perform a Cross-Document Comparative Analysis of multiple bid proposals.
{std_note}

BID DOCUMENTS:
{bids_context}

{scope_note}

Identify 5-10 key technical parameters or requirements from the documents.
For each parameter, compare what each bid proposes. If a standard is provided, also list the standard's requirement.

Return ONLY a valid JSON array of comparisons:
[
  {{
    "parameter": "E.g., Top Speed, Endurance, Hull Material",
    "standard_requirement": "What the standard requires (if available, else 'N/A')",
    "bids": [
      {{ "bidder": "Bid_Filename.pdf", "value": "What they proposed", "compliant": true/false }}
    ],
    "analysis": "Brief analysis of who is better or if someone is non-compliant",
    "winner": "Name of best bidder for this parameter or 'Tie'"
  }}
]
"""

    messages = [
        {"role": "system", "content": "You are an expert bid evaluator. Return only valid JSON arrays."},
        {"role": "user", "content": prompt},
    ]

    raw_resp = llm_engine.generate(messages, max_tokens=4096, temperature=0.2)

    try:
        cleaned = _clean_json(raw_resp)
        start = cleaned.find("[")
        end = cleaned.rfind("]") + 1
        if start == -1 or end == 0:
            raise ValueError("No JSON array found")
        comparisons = json.loads(cleaned[start:end])
    except Exception as e:
        logger.error("Failed to parse compare JSON: %s\nRaw: %s", e, raw_resp[:500])
        raise HTTPException(status_code=500, detail="Failed to parse comparison analysis.")

    job_id = str(uuid.uuid4())

    def event_stream():
        for i, comp in enumerate(comparisons):
            yield f"data: {json.dumps({'comparison': comp, 'index': i + 1, 'total': len(comparisons)})}\n\n"

        # Generate DOCX report
        docx_path = _OUTPUTS_DIR / f"{job_id}_compare.docx"
        doc = DocxDocument()
        doc.add_heading("Comparative Bid Analysis", level=1)
        doc.add_heading(f"Bids Evaluated: {', '.join(bids_data.keys())}", level=2)
        if body.standard_doc_id:
            doc.add_heading("Standard: Included", level=2)
        
        doc.add_paragraph("")

        for comp in comparisons:
            doc.add_heading(comp.get("parameter", "Parameter"), level=3)
            doc.add_paragraph(f"Standard Requirement: {comp.get('standard_requirement', 'N/A')}")
            for bid in comp.get("bids", []):
                doc.add_paragraph(f"• {bid.get('bidder', 'Bidder')}: {bid.get('value', 'N/A')} (Compliant: {bid.get('compliant', 'N/A')})")
            doc.add_paragraph(f"Analysis: {comp.get('analysis', 'N/A')}")
            doc.add_paragraph(f"Winner: {comp.get('winner', 'N/A')}")
            doc.add_paragraph("")

        doc.save(str(docx_path))

        yield f"data: {json.dumps({'done': True, 'download_url': f'/api/agent/download/{job_id}_compare.docx'})}\n\n"

        elapsed_ms = (time.time() - compare_start) * 1000
        log_usage(action_type="compare", module="reasoning", token=auth_tok, response_time_ms=elapsed_ms)

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─────────────────────────────────────────────────────────────────────
#  Phase D: Branch-Isolated Bid Comparison
#  Uses hierarchical metadata (bidder_key + problem_statement) so each
#  bidder's chunks are retrieved INDEPENDENTLY and never cross-contaminate.
# ─────────────────────────────────────────────────────────────────────


# Standard catalogue used for compliance evaluation
_STANDARDS_CATALOG = {
    "ISO_9001": "Quality Management Systems",
    "ISO_14001": "Environmental Management Systems",
    "ISO_27001": "Information Security Management",
    "ISO_45001": "Occupational Health & Safety",
    "IS_2026": "Indian Standard for steel structures",
    "IS_456": "Indian Standard for concrete structures",
    "IEEE_802": "IEEE networking standards",
    "MIL_STD_810": "Environmental engineering considerations",
    "MIL_STD_461": "Electromagnetic compatibility",
    "ICG_TECH_SOTR": "Indian Coast Guard Technical SOTR",
    "GENERIC_REQS": "Generic functional requirements (price, delivery, warranty)",
}


class BranchCompareRequest(BaseModel):
    bidder_keys: List[str] = Field(..., min_length=2, description="At least two bidder_key values")
    problem_statement: str = Field(..., description="Shared tender/problem reference all bids must match")
    standards: Optional[List[str]] = Field(None, description="Standard IDs to evaluate against. If None, returns a clarification prompt.")
    focus: Optional[str] = Field(None, description="Optional focus area, e.g. 'pricing', 'delivery timeline'")


def _strip_json_fences(raw: str) -> str:
    cleaned = re.sub(r"```(?:json)?\s*", "", raw or "").replace("```", "").strip()
    return cleaned


def _parse_json_object(raw: str) -> Optional[dict]:
    """Resilient JSON object parser. Returns None on failure."""
    cleaned = _strip_json_fences(raw)
    s = cleaned.find("{")
    e = cleaned.rfind("}") + 1
    if s < 0 or e <= s:
        return None
    try:
        return json.loads(cleaned[s:e])
    except (json.JSONDecodeError, ValueError):
        return None


def _retrieve_bidder_chunks(
    store,
    query: str,
    bidder_key: str,
    problem_statement: str,
    top_k: int = 12,
    user_clearance: int = 4,
) -> List[Dict[str, Any]]:
    """Branch-isolated retrieval: only chunks tagged with this bidder + tender."""
    query_emb = embedder.embed_texts([query])[0]
    return store.hybrid_search(
        query_text=query,
        query_embedding=query_emb,
        top_k=top_k,
        bidder_key=bidder_key,
        problem_statement=problem_statement,
        user_clearance=user_clearance,
    )


def _evaluate_one(
    bidder_key: str,
    standard_id: str,
    standard_desc: str,
    chunks: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Ask the LLM whether `chunks` (which represent ONE bid in isolation)
    comply with the given standard. Returns structured verdict.
    """
    if not chunks:
        return {
            "bidder_key": bidder_key,
            "standard_id": standard_id,
            "verdict": "insufficient_evidence",
            "severity": "high",
            "rationale": "No content found for this bidder under the specified tender.",
            "citations": [],
        }

    context = "\n\n".join(
        f"[{i+1}] (page {c.get('metadata', {}).get('page', '?')}) {c.get('text', '')[:900]}"
        for i, c in enumerate(chunks[:8])
    )
    prompt = (
        f"You are a compliance evaluator. Evaluate whether the bidder '{bidder_key}' "
        f"complies with the standard '{standard_id}' ({standard_desc}) based on the "
        "evidence excerpts below. Be strict, cite specific excerpt numbers, and use "
        "ONLY the provided evidence — never invent facts.\n\n"
        f"EVIDENCE (only this bidder's submission):\n{context}\n\n"
        "Return a JSON object with these keys:\n"
        '  "verdict": one of "compliant" | "partial" | "non_compliant" | "insufficient_evidence",\n'
        '  "severity": one of "low" | "medium" | "high",\n'
        '  "rationale": 2-3 sentences explaining the verdict,\n'
        '  "citations": list of excerpt numbers used (e.g. [1, 3, 5])\n\n'
        "JSON only:"
    )
    try:
        raw = llm_engine.generate(
            messages=[
                {"role": "system", "content": "You are a strict compliance evaluator. Output JSON only."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=600,
            temperature=0.1,
            response_format={"type": "json_object"},
            raw=True,
        )
        parsed = _parse_json_object(raw) or {}
    except Exception as ex:
        logger.warning("LLM eval failed for %s × %s: %s", bidder_key, standard_id, ex)
        parsed = {}

    return {
        "bidder_key": bidder_key,
        "standard_id": standard_id,
        "verdict": parsed.get("verdict") or "insufficient_evidence",
        "severity": parsed.get("severity") or "medium",
        "rationale": llm_engine.sanitize_text(parsed.get("rationale") or "No rationale produced."),
        "citations": parsed.get("citations") or [],
    }


def _build_executive_summary(
    bidder_keys: List[str],
    problem_statement: str,
    findings_by_bidder: Dict[str, List[dict]],
) -> str:
    """One LLM call to summarize the cross-bid comparison."""
    bullets = []
    for bk in bidder_keys:
        compliant = sum(1 for f in findings_by_bidder.get(bk, []) if f["verdict"] == "compliant")
        partial = sum(1 for f in findings_by_bidder.get(bk, []) if f["verdict"] == "partial")
        noncomp = sum(1 for f in findings_by_bidder.get(bk, []) if f["verdict"] == "non_compliant")
        insuff = sum(1 for f in findings_by_bidder.get(bk, []) if f["verdict"] == "insufficient_evidence")
        bullets.append(
            f"- {bk}: {compliant} compliant, {partial} partial, {noncomp} non-compliant, {insuff} insufficient"
        )
    summary_input = (
        f"Tender: {problem_statement}\nBidders: {', '.join(bidder_keys)}\n\n"
        + "\n".join(bullets)
        + "\n\nWrite a concise 4-6 sentence executive summary highlighting which bidder "
        "is stronger overall and where each has gaps. Plain prose, no bullets, no headings."
    )
    try:
        raw = llm_engine.generate(
            messages=[
                {"role": "system", "content": "You are a procurement evaluation officer."},
                {"role": "user", "content": summary_input},
            ],
            max_tokens=500,
            temperature=0.3,
            raw=False,
        )
        return llm_engine.sanitize_text(raw.strip())
    except Exception as ex:
        logger.warning("Executive summary generation failed: %s", ex)
        return "Executive summary unavailable. See per-standard findings below."


def _build_recommendation(
    bidder_keys: List[str],
    findings_by_bidder: Dict[str, List[dict]],
) -> Dict[str, Any]:
    """Pick the recommended bidder based on weighted verdict counts."""
    scores: Dict[str, float] = {}
    weights = {"compliant": 1.0, "partial": 0.5, "non_compliant": -0.5, "insufficient_evidence": 0.0}
    for bk in bidder_keys:
        s = 0.0
        for f in findings_by_bidder.get(bk, []):
            s += weights.get(f["verdict"], 0.0)
        scores[bk] = round(s, 2)

    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    if not ranked:
        return {"recommended": None, "scores": scores, "reasoning": "No findings to evaluate."}

    top_bk, top_score = ranked[0]
    if len(ranked) > 1 and abs(ranked[0][1] - ranked[1][1]) < 0.5:
        return {
            "recommended": None,
            "scores": scores,
            "reasoning": f"Bidders are within 0.5 points — manual review recommended.",
        }
    return {
        "recommended": top_bk,
        "scores": scores,
        "reasoning": f"{top_bk} scored {top_score} vs runner-up {ranked[1][0] if len(ranked) > 1 else 'n/a'}.",
    }


def _generate_compare_docx(
    job_id: str,
    bidder_keys: List[str],
    problem_statement: str,
    standards_used: List[str],
    findings_by_bidder: Dict[str, List[dict]],
    exec_summary: str,
    recommendation: Dict[str, Any],
) -> Path:
    docx_path = _OUTPUTS_DIR / f"{job_id}_branch_compare.docx"
    doc = DocxDocument()
    doc.add_heading("Branch-Isolated Bid Comparison Report", level=1)
    doc.add_paragraph(f"Tender / Problem Statement: {problem_statement}")
    doc.add_paragraph(f"Bidders Evaluated: {', '.join(bidder_keys)}")
    doc.add_paragraph(f"Standards Applied: {', '.join(standards_used)}")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(exec_summary)

    doc.add_heading("Recommendation", level=2)
    rec_text = (
        f"Recommended: {recommendation.get('recommended') or 'Tie / manual review'}\n"
        f"Reasoning: {recommendation.get('reasoning', '')}\n"
        f"Scores: {recommendation.get('scores', {})}"
    )
    doc.add_paragraph(rec_text)

    doc.add_heading("Per-Standard Findings", level=2)
    for std in standards_used:
        doc.add_heading(f"{std} — {_STANDARDS_CATALOG.get(std, '')}", level=3)
        for bk in bidder_keys:
            finding = next((f for f in findings_by_bidder.get(bk, []) if f["standard_id"] == std), None)
            if not finding:
                continue
            p = doc.add_paragraph()
            p.add_run(f"{bk}: ").bold = True
            p.add_run(f"{finding['verdict']} ({finding['severity']} severity)\n")
            p.add_run(finding["rationale"])
        doc.add_paragraph("")

    # Watermark
    try:
        footer = doc.sections[0].footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = "AI-Generated Draft — ICG AGRA — Branch-Isolated Comparison"
    except Exception:
        pass

    doc.save(str(docx_path))
    return docx_path


@router.get("/compare/standards/catalog")
async def compare_standards_catalog(user: dict = Depends(get_current_user)):
    """Return the catalog of standards the comparison engine can evaluate."""
    return {
        "standards": [
            {"id": sid, "description": desc} for sid, desc in _STANDARDS_CATALOG.items()
        ]
    }


@router.get("/compare/bids/available")
async def compare_bids_available(user: dict = Depends(get_current_user)):
    """
    Return the catalog of bidders and problem_statements currently indexed,
    grouped so the UI can offer valid comparison combinations.
    """
    store = get_store()
    data = store.list_bid_documents()
    # Annotate which problem statements have enough bidders to be comparable (>= 2)
    for ps in data["problem_statements"]:
        ps["comparable"] = len(ps["bidder_keys"]) >= 2
    return data


@router.post("/compare/bids/branch")
async def compare_bids_branch(
    body: BranchCompareRequest,
    request: Request,
    user: dict = Depends(get_current_user),
):
    """
    Phase D — Branch-isolated bid comparison.

    Workflow:
      1. Verify each bidder_key has chunks tagged with the shared problem_statement.
      2. Retrieve chunks per bidder INDEPENDENTLY (no cross-contamination).
      3. Run compliance evaluation per (bidder, standard) pair.
      4. Aggregate, generate exec summary + recommendation + DOCX.
      5. Stream progress via SSE.
    """
    compare_start = time.time()
    auth_header = request.headers.get("authorization", "") if request else ""
    auth_tok = auth_header.replace("Bearer ", "") if auth_header else ""

    # ── If standards not provided, return a clarification prompt ──
    if not body.standards:
        catalog = [
            {"id": sid, "description": desc}
            for sid, desc in _STANDARDS_CATALOG.items()
        ]
        return {
            "needs_clarification": True,
            "message": (
                "Which standards would you like the bids evaluated against? "
                "Select one or more from the catalog and resubmit."
            ),
            "catalog": catalog,
        }

    # ── Validate standards ──
    unknown = [s for s in body.standards if s not in _STANDARDS_CATALOG]
    if unknown:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown standard IDs: {unknown}. Call GET /compare/standards/catalog for valid IDs.",
        )

    store = get_store()
    bidder_keys = body.bidder_keys
    problem_statement = body.problem_statement
    standards = body.standards
    focus = body.focus or ""
    user_clearance = int(user.get("clearance_level", 4))

    # ── Cross-bid validation: each bidder must have at least one chunk ──
    # under the same problem_statement.
    for bk in bidder_keys:
        probe = store.hybrid_search(
            query_text=focus or problem_statement,
            query_embedding=embedder.embed_texts([focus or problem_statement])[0],
            top_k=1,
            bidder_key=bk,
            problem_statement=problem_statement,
            user_clearance=user_clearance,
        )
        if not probe:
            raise HTTPException(
                status_code=422,
                detail=(
                    f"No indexed content found for bidder '{bk}' under tender "
                    f"'{problem_statement}'. Make sure both bids were uploaded with "
                    f"matching problem_statement and the bidder_key spelled correctly."
                ),
            )

    job_id = str(uuid.uuid4())

    async def event_stream():
        total_steps = len(bidder_keys) * len(standards)
        step = 0
        findings_by_bidder: Dict[str, List[dict]] = {bk: [] for bk in bidder_keys}

        yield f"data: {json.dumps({'stage': 'started', 'job_id': job_id, 'total_steps': total_steps})}\n\n"

        for std_id in standards:
            std_desc = _STANDARDS_CATALOG[std_id]
            # Query crafted to elicit comparable evidence per bidder
            std_query = f"{std_desc}. {focus}".strip()
            for bk in bidder_keys:
                step += 1
                yield (
                    "data: " + json.dumps({
                        "stage": "evaluating",
                        "bidder_key": bk,
                        "standard_id": std_id,
                        "progress": int(step * 100 / max(1, total_steps)),
                        "step": step,
                        "total_steps": total_steps,
                    }) + "\n\n"
                )
                # Branch-isolated retrieval
                bid_chunks = await asyncio.to_thread(
                    _retrieve_bidder_chunks, store, std_query, bk,
                    problem_statement, 10, user_clearance,
                )
                finding = await asyncio.to_thread(
                    _evaluate_one, bk, std_id, std_desc, bid_chunks,
                )
                findings_by_bidder[bk].append(finding)
                yield f"data: {json.dumps({'stage': 'finding', 'finding': finding})}\n\n"

        # ── Aggregate ──
        yield f"data: {json.dumps({'stage': 'aggregating', 'progress': 95})}\n\n"
        exec_summary = await asyncio.to_thread(
            _build_executive_summary, bidder_keys, problem_statement, findings_by_bidder,
        )
        recommendation = _build_recommendation(bidder_keys, findings_by_bidder)

        # Build standards table (cross-bidder per standard)
        standards_table = []
        for std_id in standards:
            row: Dict[str, Any] = {
                "standard_id": std_id,
                "description": _STANDARDS_CATALOG[std_id],
                "per_bidder": {},
            }
            for bk in bidder_keys:
                f = next((x for x in findings_by_bidder[bk] if x["standard_id"] == std_id), None)
                row["per_bidder"][bk] = {
                    "verdict": f["verdict"] if f else "insufficient_evidence",
                    "severity": f["severity"] if f else "medium",
                } if f else None
            standards_table.append(row)

        # ── DOCX ──
        try:
            docx_path = await asyncio.to_thread(
                _generate_compare_docx,
                job_id, bidder_keys, problem_statement, standards,
                findings_by_bidder, exec_summary, recommendation,
            )
            download_url = f"/api/agent/download/{docx_path.name}"
        except Exception as ex:
            logger.exception("DOCX generation failed: %s", ex)
            download_url = None

        elapsed_ms = round((time.time() - compare_start) * 1000, 1)
        log_usage(
            action_type="compare_branch",
            module="reasoning",
            token=auth_tok,
            response_time_ms=elapsed_ms,
        )

        yield (
            "data: " + json.dumps({
                "stage": "done",
                "job_id": job_id,
                "executive_summary": exec_summary,
                "recommendation": recommendation,
                "standards_table": standards_table,
                "findings_by_bidder": findings_by_bidder,
                "bidder_keys": bidder_keys,
                "problem_statement": problem_statement,
                "standards_used": standards,
                "elapsed_ms": elapsed_ms,
                "download_url": download_url,
            }) + "\n\n"
        )

    return StreamingResponse(event_stream(), media_type="text/event-stream")
