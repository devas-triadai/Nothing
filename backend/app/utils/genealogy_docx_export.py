"""
Module 7 Phase 5 — Genealogy DOCX Export
Generate professional Word documents with document lineage tables.
"""

import logging
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sqlalchemy.orm import Session
from app.models.models import (
    Document, DocEdge, ExtractedDocumentMetadata, 
    DocumentEntity, DocumentChangeSummary, User
)

logger = logging.getLogger(__name__)


def _add_watermark(doc: Document, text: str = "AI-Generated Report - ICG AGRA"):
    """Add watermark text to document header."""
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Make watermark gray and small
        run = header_para.runs[0] if header_para.runs else header_para.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)


def _add_classification_banner(doc: Document, classification: str = "UNCLASSIFIED"):
    """Add classification banner to top and bottom of each page."""
    banner_text = f"CLASSIFICATION: {classification}"
    
    for section in doc.sections:
        # Top banner
        header = section.header
        top_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        top_para.text = banner_text
        top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = top_para.runs[0] if top_para.runs else top_para.add_run(banner_text)
        run.font.bold = True
        run.font.size = Pt(10)
        
        # Color based on classification
        if classification in ["CONFIDENTIAL", "SECRET", "TOP SECRET"]:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for classified
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green for unclassified
        
        # Bottom banner
        footer = section.footer
        bottom_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        bottom_para.text = banner_text
        bottom_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bottom_para.runs[0] if bottom_para.runs else bottom_para.add_run(banner_text)
        run.font.bold = True
        run.font.size = Pt(10)
        
        if classification in ["CONFIDENTIAL", "SECRET", "TOP SECRET"]:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)


def _set_cell_shading(cell, color: str):
    """Set background color for table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def generate_genealogy_docx(
    doc_id: int,
    db: Session,
    include_entities: bool = True,
    include_changes: bool = True,
    classification: str = "UNCLASSIFIED"
) -> str:
    """
    Generate a comprehensive genealogy report as DOCX.
    
    Args:
        doc_id: Document ID to generate report for
        db: Database session
        include_entities: Include entity summary table
        include_changes: Include change summaries
        classification: Document classification level
    
    Returns:
        Path to generated DOCX file
    """
    # Get document
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise ValueError(f"Document {doc_id} not found")
    
    # Create document
    document = Document()
    
    # Set up section for headers/footers
    section = document.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    
    # Add watermark and classification
    _add_watermark(document)
    if classification != "UNCLASSIFIED":
        _add_classification_banner(document, classification)
    
    # TITLE PAGE
    document.add_heading("DOCUMENT GENEALOGY & LINEAGE REPORT", level=0)
    document.add_heading(f"{doc.original_filename}", level=1)
    
    # Metadata summary
    p = document.add_paragraph()
    p.add_run(f"Document ID: ").bold = True
    p.add_run(f"{doc.id}\n")
    p.add_run(f"Current Version: ").bold = True
    p.add_run(f"v{doc.version}\n")
    p.add_run(f"Status: ").bold = True
    p.add_run(f"{doc.status}\n")
    p.add_run(f"Category: ").bold = True
    p.add_run(f"{doc.category or 'Unclassified'}\n")
    p.add_run(f"Generated: ").bold = True
    p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p.add_run(f"SHA256: ").bold = True
    p.add_run(f"{doc.sha256_hash[:32]}..." if doc.sha256_hash else "N/A")
    
    document.add_paragraph()
    
    # SECTION 1: VERSION HISTORY TABLE
    document.add_heading("1. Document Version History", level=1)
    
    # Get all versions in this document group
    if doc.doc_group_id:
        all_versions = db.query(Document).filter(
            Document.doc_group_id == doc.doc_group_id
        ).order_by(Document.version.asc()).all()
    else:
        all_versions = [doc]
    
    if len(all_versions) > 1:
        # Create version history table
        table = document.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Version', 'Date', 'Status', 'Author', 'Changes']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            _set_cell_shading(hdr_cells[i], 'D9E2F3')  # Light blue header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for version_doc in all_versions:
            row_cells = table.add_row().cells
            row_cells[0].text = f"v{version_doc.version}"
            row_cells[1].text = version_doc.created_at.strftime('%Y-%m-%d') if version_doc.created_at else 'N/A'
            row_cells[2].text = version_doc.status or 'Unknown'
            
            # Get uploader name
            uploader = db.query(User).filter(User.id == version_doc.uploaded_by).first()
            row_cells[3].text = uploader.username if uploader else 'Unknown'
            
            # Version notes
            notes = version_doc.version_notes or 'No notes'
            if len(notes) > 100:
                notes = notes[:97] + '...'
            row_cells[4].text = notes
        
        document.add_paragraph()
    else:
        p = document.add_paragraph("This document has no version history (single version).")
        p.italic = True
        document.add_paragraph()
    
    # SECTION 2: RELATIONSHIP GRAPH TABLE
    document.add_heading("2. Document Relationships", level=1)
    
    # Get explicit edges
    outgoing = db.query(DocEdge, Document).join(
        Document, DocEdge.target_id == Document.id
    ).filter(DocEdge.source_id == doc_id).all()
    
    incoming = db.query(DocEdge, Document).join(
        Document, DocEdge.source_id == Document.id
    ).filter(DocEdge.target_id == doc_id).all()
    
    if outgoing or incoming:
        table = document.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        headers = ['Relationship', 'Target Document', 'Type', 'Confidence']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            _set_cell_shading(hdr_cells[i], 'D9E2F3')
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Outgoing relationships (this doc -> others)
        for edge, target_doc in outgoing:
            row_cells = table.add_row().cells
            row_cells[0].text = "-> Outgoing"
            row_cells[1].text = target_doc.original_filename
            row_cells[2].text = edge.edge_type
            row_cells[3].text = f"{edge.confidence:.2f}"
        
        # Incoming relationships (others -> this doc)
        for edge, source_doc in incoming:
            row_cells = table.add_row().cells
            row_cells[0].text = "<- Incoming"
            row_cells[1].text = source_doc.original_filename
            row_cells[2].text = edge.edge_type
            row_cells[3].text = f"{edge.confidence:.2f}"
        
        document.add_paragraph()
    else:
        p = document.add_paragraph("No explicit document relationships recorded.")
        p.italic = True
        document.add_paragraph()
    
    # SECTION 3: EXTRACTED METADATA
    extracted_meta = db.query(ExtractedDocumentMetadata).filter(
        ExtractedDocumentMetadata.document_id == doc_id
    ).first()
    
    if extracted_meta:
        document.add_heading("3. Extracted Metadata", level=1)
        
        # Version references
        if extracted_meta.version_refs:
            p = document.add_paragraph()
            p.add_run("Version References: ").bold = True
            p.add_run(", ".join(extracted_meta.version_refs))
        
        # Cross references
        if extracted_meta.cross_references:
            p = document.add_paragraph()
            p.add_run("Cross References: ").bold = True
            refs = [f"{r.get('doc', 'Unknown')} {r.get('ref', '')}" 
                   for r in extracted_meta.cross_references[:5]]
            p.add_run("; ".join(refs))
        
        # Amendment dates
        if extracted_meta.amendment_dates:
            p = document.add_paragraph()
            p.add_run("Amendment Dates: ").bold = True
            p.add_run(", ".join(extracted_meta.amendment_dates))
        
        # Effective date
        if extracted_meta.effective_date:
            p = document.add_paragraph()
            p.add_run("Effective Date: ").bold = True
            p.add_run(extracted_meta.effective_date.strftime('%Y-%m-%d'))
        
        # Supersession info
        if extracted_meta.supersession_info:
            p = document.add_paragraph()
            p.add_run("Supersession: ").bold = True
            info = extracted_meta.supersession_info
            parts = []
            if info.get('supersedes'):
                parts.append(f"Supersedes: {info['supersedes']}")
            if info.get('superseded_by'):
                parts.append(f"Superseded by: {info['superseded_by']}")
            p.add_run("; ".join(parts))
        
        document.add_paragraph()
    
    # SECTION 4: ENTITY SUMMARY
    if include_entities:
        entities = db.query(DocumentEntity).filter(
            DocumentEntity.document_id == doc_id
        ).all()
        
        if entities:
            document.add_heading("4. Extracted Entities", level=1)
            
            # Group by type
            by_type: Dict[str, List] = {}
            for entity in entities:
                if entity.entity_type not in by_type:
                    by_type[entity.entity_type] = []
                by_type[entity.entity_type].append(entity)
            
            # Summary table
            table = document.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            headers = ['Entity Type', 'Count', 'Examples']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                _set_cell_shading(hdr_cells[i], 'D9E2F3')
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for entity_type, entity_list in sorted(by_type.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = entity_type.replace('_', ' ').title()
                row_cells[1].text = str(len(entity_list))
                
                # Examples (top 3 by confidence)
                top_examples = sorted(entity_list, 
                                    key=lambda e: e.extraction_confidence, 
                                    reverse=True)[:3]
                examples = [e.entity_name for e in top_examples]
                row_cells[2].text = ", ".join(examples)
            
            document.add_paragraph()
    
    # SECTION 5: CHANGE SUMMARIES
    if include_changes:
        changes = db.query(DocumentChangeSummary).filter(
            (DocumentChangeSummary.from_doc_id == doc_id) | 
            (DocumentChangeSummary.to_doc_id == doc_id)
        ).all()
        
        if changes:
            document.add_heading("5. Version Change Summaries", level=1)
            
            for change in changes:
                direction = "to" if change.from_doc_id == doc_id else "from"
                other_id = change.to_doc_id if change.from_doc_id == doc_id else change.from_doc_id
                other_doc = db.query(Document).filter(Document.id == other_id).first()
                
                p = document.add_paragraph()
                other_ver = other_doc.version if other_doc else '?'
                run = p.add_run(f"Changes {direction} v{other_ver}:")
                run.bold = True
                p.add_run(f" [{change.impact_assessment or 'Unknown'} Impact]")
                
                if change.summary_text:
                    document.add_paragraph(change.summary_text, style='List Bullet')
                
                if change.major_changes:
                    document.add_paragraph("Major Changes:", style='List Bullet')
                    for item in change.major_changes[:5]:
                        document.add_paragraph(f"  - {item}", style='List Bullet 2')
                
                document.add_paragraph()
    
    # SECTION 6: APPENDIX - INTEGRITY
    document.add_page_break()
    document.add_heading("Appendix: Document Integrity", level=1)
    
    p = document.add_paragraph()
    p.add_run("SHA256 Hashes:\n").bold = True
    
    if all_versions:
        for v in all_versions:
            p.add_run(f"v{v.version}: ")
            p.add_run(f"{v.sha256_hash or 'N/A'}\n")
    else:
        p.add_run(f"v{doc.version}: {doc.sha256_hash or 'N/A'}\n")
    
    # Generation metadata
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Report Metadata:\n").bold = True
    p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p.add_run(f"System: ICG AGRA Document Management System\n")
    p.add_run(f"Classification: {classification}\n")
    
    # Save document
    output_dir = Path("/tmp/agra_exports")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"genealogy_doc_{doc_id}_{timestamp}.docx"
    output_path = output_dir / filename
    
    document.save(str(output_path))
    
    logger.info("Generated genealogy report: %s", output_path)
    
    return str(output_path)
